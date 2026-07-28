# -*- coding: utf-8 -*-
"""Report template export system - placeholder fill, chart gen, table fill, summary gen."""
import os
import io
import json
import re
import traceback
import copy
import pandas as pd
import requests
from flask import request, jsonify, send_file
from sqlalchemy import text as sa_text

try:
    from common import protected as _protected
except ImportError:
    from helpers import protected as _protected

DOUBAO_API_KEY = os.getenv('DOUBAO_API_KEY', '')
DOUBAO_API_URL = os.getenv('DOUBAO_API_URL', 'https://ark.cn-beijing.volces.com/api/v3/chat/completions')
DOUBAO_MODEL = os.getenv('DOUBAO_MODEL', 'doubao-seed-1-8-251228')

CREATE_TABLE_SQL = """
CREATE TABLE IF NOT EXISTS report_templates (
    id INT AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(200) NOT NULL,
    description VARCHAR(500) DEFAULT '',
    report_type VARCHAR(20) NOT NULL DEFAULT 'single',
    sections JSON NOT NULL,
    template_file VARCHAR(500) DEFAULT NULL,
    template_structure JSON DEFAULT NULL,
    created_by INT NOT NULL DEFAULT 0,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;
"""

ALTER_TABLE_SQL = [
    "ALTER TABLE report_templates ADD COLUMN template_file VARCHAR(500) DEFAULT NULL",
    "ALTER TABLE report_templates ADD COLUMN template_structure JSON DEFAULT NULL",
]

SECTION_QUERY_MAP = {
    "案件总览": "统计案件总量、每日趋势、类型分布",
    "处置效率": "统计平均处置时长、办结率、时长分布",
    "问题大类分布": "统计各问题大类的案件数量和占比",
    "小类 Top": "统计案件数量最多的前15个小类",
    "处置部门分析": "统计各处置部门的案件数量和平均时长",
    "延期与返工": "统计延期和返工案件的数量及部门分布",
    "重复案件分析": "统计重复出现的案件(问题描述+小类均相近即为重复)",
}

SECTION_CHART_CONFIG = {
    "案件总览": [
        {"title": "每日上报趋势", "chart_type": "line", "query": "SELECT DATE(report_time) AS 日期, COUNT(*) AS 案件数 FROM case_data {time_filter} GROUP BY DATE(report_time) ORDER BY 日期"},
        {"title": "24小时分布", "chart_type": "bar", "query": "SELECT HOUR(report_time) AS 小时, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY HOUR(report_time) ORDER BY 小时"},
        {"title": "问题来源占比", "chart_type": "pie", "query": "SELECT source AS 来源, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY source ORDER BY 案件数 DESC"},
    ],
    "处置效率": [
        {"title": "办结率", "chart_type": "completion_rate", "query": "SELECT CASE WHEN close_time IS NOT NULL THEN '已办结' ELSE '未办结' END AS 状态, COUNT(*) AS 数量, ROUND(COUNT(*)/(SELECT COUNT(*) FROM case_data {month_filter})*100,1) AS 百分比 FROM case_data {month_filter} GROUP BY 状态"},
        {"title": "处置时长分布", "chart_type": "duration_dist", "query": "SELECT CASE WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=4 THEN '0-4h' WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=8 THEN '4-8h' WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=12 THEN '8-12h' WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=24 THEN '12-24h' WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=48 THEN '1-2天' WHEN TIMESTAMPDIFF(HOUR, report_time, close_time)<=96 THEN '2-4天' ELSE '>4天' END AS 时段, COUNT(*) AS 案件数, ROUND(COUNT(*)/(SELECT COUNT(*) FROM case_data WHERE close_time IS NOT NULL {month_filter_and})*100,1) AS 百分比 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY 时段 ORDER BY FIELD(时段,'0-4h','4-8h','8-12h','12-24h','1-2天','2-4天','>4天')"},
    ],
    "问题大类分布": [
        {"title": "大类案件数量", "chart_type": "bar", "query": "SELECT big_category AS 大类, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY big_category ORDER BY 案件数 DESC"},
        {"title": "大类占比", "chart_type": "pie", "query": "SELECT big_category AS 大类, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY big_category ORDER BY 案件数 DESC"},
    ],
    "小类 Top": [
        {"title": "小类案件排名 Top15", "chart_type": "horizontal_bar", "query": "SELECT small_category AS 小类, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY small_category ORDER BY 案件数 DESC LIMIT 15"},
    ],
    "处置部门分析": [
        {"title": "部门案件数量", "chart_type": "bar", "query": "SELECT department AS 部门, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY department ORDER BY 案件数 DESC LIMIT 15"},
        {"title": "部门平均处置时长", "chart_type": "horizontal_bar", "query": "SELECT department AS 部门, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,2) AS 平均时长 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY department ORDER BY 平均时长 DESC LIMIT 15"},
    ],
    "延期与返工": [
        {"title": "延期案件部门分布", "chart_type": "bar", "query": "SELECT department AS 部门, COUNT(*) AS 延期数 FROM case_data {month_filter} AND is_delayed=1 GROUP BY department ORDER BY 延期数 DESC"},
        {"title": "返工案件部门分布", "chart_type": "bar", "query": "SELECT department AS 部门, COUNT(*) AS 返工数 FROM case_data {month_filter} AND is_rework=1 GROUP BY department ORDER BY 返工数 DESC"},
    ],
    "重复案件分析": [
        {"title": "重复次数分布", "chart_type": "bar", "query": "SELECT repeat_cnt AS 次数, COUNT(*) AS 组数 FROM (SELECT CONCAT(LEFT(description,30),'-',small_category) AS case_key, COUNT(*) AS repeat_cnt FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY case_key HAVING repeat_cnt>1) t GROUP BY repeat_cnt ORDER BY repeat_cnt"},
        {"title": "顽固重复店铺 Top12", "chart_type": "horizontal_bar", "query": "SELECT LEFT(description,40) AS 地点, COUNT(*) AS 重复次数 FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY LEFT(description,40) HAVING 重复次数>1 ORDER BY 重复次数 DESC LIMIT 12"},
    ],
}

# Comparison report chart configurations (for 2-month comparison)
SECTION_CHART_CONFIG_COMPARE = {
    "案件数量对比": [
        {
            "title": "案件总量对比",
            "chart_type": "compare_total",
            "query": "SELECT upload_batch AS 月份, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY upload_batch ORDER BY 月份"
        },
    ],
    "处置效率对比": [
        {
            "title": "平均处置时长与办结率",
            "chart_type": "compare_efficiency",
            "queries": [
                "SELECT upload_batch AS 月份, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY upload_batch ORDER BY 月份",
                "SELECT upload_batch AS 月份, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),1) AS 办结率 FROM case_data {month_filter} GROUP BY upload_batch ORDER BY 月份"
            ]
        },
    ],
    "延期与返工": [
        {
            "title": "延期与返工数量及比率",
            "chart_type": "compare_delay_rework",
            "queries": [
                "SELECT upload_batch AS 月份, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {month_filter} GROUP BY upload_batch ORDER BY 月份",
                "SELECT upload_batch AS 月份, ROUND(SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),2) AS 延期率, ROUND(SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),2) AS 返工率 FROM case_data {month_filter} GROUP BY upload_batch ORDER BY 月份"
            ]
        },
    ],
    "问题大类变化": [
        {
            "title": "大类数量对比及变化率",
            "chart_type": "compare_category_change",
            "query": "SELECT big_category AS 大类, upload_batch AS 月份, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY big_category, upload_batch ORDER BY 案件数 DESC"
        },
    ],
    "小类 Top": [
        {
            "title": "小类案件数量变化",
            "chart_type": "compare_small_category",
            "query": "SELECT small_category AS 小类, upload_batch AS 月份, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY small_category, upload_batch ORDER BY 案件数 DESC"
        },
    ],
    "重复案件分析": [
        {
            "title": "重复案件状态及顽固案件",
            "chart_type": "compare_repeat",
            "queries": [
                "SELECT 状态, 数量 FROM (SELECT '持续存在' AS 状态, COUNT(*) AS 数量 FROM (SELECT street, LEFT(description,40) AS description, big_category, small_category FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY street, LEFT(description,40), big_category, small_category HAVING COUNT(DISTINCT upload_batch) = 2) t UNION ALL SELECT '新增' AS 状态, COUNT(*) AS 数量 FROM (SELECT street, LEFT(description,40) AS description, big_category, small_category FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY street, LEFT(description,40), big_category, small_category HAVING COUNT(DISTINCT upload_batch) = 1 AND MIN(upload_batch) = '{m2}') t UNION ALL SELECT '已解决' AS 状态, (SELECT COUNT(*) FROM (SELECT street, LEFT(description,40) AS description, big_category, small_category FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY street, LEFT(description,40), big_category, small_category HAVING COUNT(DISTINCT upload_batch) = 1 AND MAX(upload_batch) = '{m1}') t2)) t3 ORDER BY FIELD(状态,'持续存在','已解决','新增')",
                "SELECT CONCAT(店铺,'_',LEFT(描述,20)) AS 店铺, 月份, 次数 FROM (SELECT street AS 店铺, description AS 描述, upload_batch AS 月份, COUNT(*) AS 次数 FROM case_data {month_filter} AND description IS NOT NULL AND description!='' GROUP BY street, description, upload_batch HAVING COUNT(*) > 1) t ORDER BY 次数 DESC LIMIT 15"
            ]
        },
    ],
    "处置部门分析": [
        {
            "title": "部门平均处置时长对比",
            "chart_type": "compare_department",
            "query": "SELECT department AS 部门, upload_batch AS 月份, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, COUNT(*) AS 案件数 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY department, upload_batch ORDER BY 案件数 DESC"
        },
    ],
    "处置部门效率": [
        {
            "title": "部门平均处置时长对比",
            "chart_type": "compare_department",
            "query": "SELECT department AS 部门, upload_batch AS 月份, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, COUNT(*) AS 案件数 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY department, upload_batch ORDER BY 案件数 DESC"
        },
    ],
    "问题来源变化": [
        {
            "title": "问题来源渠道对比",
            "chart_type": "compare_source_change",
            "query": "SELECT source AS 来源, upload_batch AS 月份, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY source, upload_batch ORDER BY 案件数 DESC"
        },
    ],
}


def _build_month_filter(selected_months):
    """Build upload_batch WHERE condition from selected month strings like ['202606']"""
    if not selected_months:
        return ""
    conditions = []
    for m in selected_months:
        if len(m) >= 6:
            conditions.append(f"upload_batch='{m[:6]}'")
    if not conditions:
        return ""
    if len(conditions) == 1:
        return "WHERE " + conditions[0]
    return "WHERE (" + " OR ".join(conditions) + ")"


# ============================================================
# Filter builders
# ============================================================

def _build_time_filter(selected_months):
    """Build SQL filter by report_time in selected months (for daily trend only)"""
    if not selected_months:
        return ""
    conditions = []
    for m in selected_months:
        if len(m) >= 6:
            year, month = m[:4], m[4:6]
            conditions.append(f"(YEAR(report_time)={year} AND MONTH(report_time)={month})")
    if not conditions:
        return ""
    if len(conditions) == 1:
        return "WHERE " + conditions[0]
    return "WHERE (" + " OR ".join(conditions) + ")"


def _get_chart_configs_with_month(section_title, selected_months, report_type='single'):
    """Get chart configs with month filter applied to all queries."""
    import re
    config_source = SECTION_CHART_CONFIG_COMPARE if report_type == 'compare' else SECTION_CHART_CONFIG
    configs = config_source.get(section_title, None)
    if configs is None:
        clean_title = re.sub(r'^[一二三四五六七八九十]+[、．.]\s*', '', section_title)
        configs = config_source.get(clean_title, None)
    if configs is None:
        for key in config_source:
            if key in section_title:
                configs = config_source[key]
                break
    if configs is None:
        clean_title = re.sub(r'^[一二三四五六七八九十]+[、．.]\s*', '', section_title)
        for key in config_source:
            if sorted(key) == sorted(clean_title):
                configs = config_source[key]
                break
    if configs is None:
        configs = []

    month_filter = _build_month_filter(selected_months)
    month_filter_and = month_filter.replace('WHERE', 'AND', 1) if month_filter.startswith('WHERE') else month_filter
    time_filter = _build_time_filter(selected_months)
    time_filter_and = time_filter.replace('WHERE', 'AND', 1) if time_filter.startswith('WHERE') else time_filter

    result = []
    for cfg in configs:
        title = cfg["title"]
        chart_type = cfg.get("chart_type", "bar")

        def _substitute(q):
            m1_val = selected_months[0] if len(selected_months) >= 1 else ''
            m2_val = selected_months[1] if len(selected_months) >= 2 else ''
            q = q.replace("{m1}", m1_val)
            q = q.replace("{m2}", m2_val)
            q = q.replace("{month_filter_and}", "\x00MF_AND\x00")
            q = q.replace("{month_filter}", month_filter if month_filter else "\x00MF_EMPTY\x00")
            q = q.replace("\x00MF_AND\x00", month_filter_and)
            q = q.replace("\x00MF_EMPTY\x00", "")
            if not month_filter:
                q = re.sub(r'\bFROM\s+case_data\s+AND\b', 'FROM case_data', q, flags=re.IGNORECASE)
            return q

        if "queries" in cfg:
            processed_queries = [_substitute(q) for q in cfg["queries"]]
            entry = {"title": title, "chart_type": chart_type, "queries": processed_queries}
            if "query" in cfg:
                entry["query"] = _substitute(cfg["query"])
            result.append(entry)
        else:
            query = _substitute(cfg["query"])
            if "每日上报趋势" in title:
                if re.search(r'WHERE\s+\S+.*\{time_filter\}', query, re.IGNORECASE | re.DOTALL):
                    query = query.replace("{time_filter}", time_filter_and)
                else:
                    query = query.replace("{time_filter}", time_filter)
                query = query.replace("{time_filter}", "")
            result.append({"title": title, "chart_type": chart_type, "query": query})
    return result


def infer_query_from_title(title):
    """Infer analysis query from section title"""
    for keyword, query in SECTION_QUERY_MAP.items():
        if keyword in title:
            return query
    return f"分析{title}相关数据"


# ============================================================
# LLM helper
# ============================================================

def _call_llm(messages, timeout=120, max_retries=3):
    """Call Doubao API with retry"""
    import time

    formatted_messages = []
    for msg in messages:
        content = msg.get("content", "")
        if isinstance(content, str):
            formatted_messages.append({
                "role": msg["role"],
                "content": [{"type": "text", "text": content}]
            })
        else:
            formatted_messages.append(msg)

    for attempt in range(max_retries):
        try:
            response = requests.post(
                DOUBAO_API_URL,
                headers={"Content-Type": "application/json", "Authorization": f"Bearer {DOUBAO_API_KEY}"},
                json={"model": DOUBAO_MODEL, "messages": formatted_messages, "temperature": 0.1, "max_tokens": 2000},
                timeout=timeout
            )
            if response.status_code == 200:
                data = response.json()
                choices = data.get("choices", [])
                if choices and choices[0].get("message", {}).get("content"):
                    return choices[0]["message"]["content"]
                return None
            else:
                print(f"[TemplateExport] API error (attempt {attempt+1}/{max_retries}): HTTP {response.status_code}")
                if attempt < max_retries - 1:
                    time.sleep(2 * (attempt + 1))
        except Exception as e:
            print(f"[TemplateExport] LLM exception: {e}")
            if attempt < max_retries - 1:
                time.sleep(2 * (attempt + 1))
    return None

# ============================================================
# Summary data
# ============================================================

def _get_summary_data(engine, selected_months):
    """Get summary data filtered by selected months"""
    month_filter = _build_month_filter(selected_months)
    sql = f"SELECT COUNT(*) AS 案件总量, CONCAT(ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1),'h') AS 平均处置时长, CONCAT(ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)/COUNT(*)*100,1),'%') AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {month_filter}"
    try:
        with engine.connect() as conn:
            df = pd.read_sql(sa_text(sql), conn)
            if not df.empty:
                result = df.iloc[0].to_dict()
                for key in ["案件总量", "延期案件", "返工案件"]:
                    if key in result and isinstance(result[key], float):
                        result[key] = int(result[key])
                return result
    except Exception as e:
        print(f"获取汇总数据失败: {e}")
    return {"案件总量": "-", "平均处置时长": "-", "办结率": "-", "延期案件": "-", "返工案件": "-"}


def _extract_shop_from_desc(desc):
    """从问题描述中提取店铺/地点名称
    格式: "街道名，店铺名，问题类型"
    例如: "西城墙路，火锅涮牛肚，店外经营。" -> "火锅涮牛肚"
    """
    if not desc or not isinstance(desc, str):
        return ''
    parts = re.split(r'[，,]', desc)
    if len(parts) < 2:
        return ''
    shop = parts[1].strip()
    shop = re.sub(r'^(与|在|从|到)', '', shop)
    shop = re.sub(r'(前|后|东侧|西侧|南侧|北侧|路口|门口|对面|斜对面|旁边|附近|之上|之下).*$', '', shop)
    shop = re.sub(r'^的', '', shop)
    shop = shop.strip()
    return shop if 2 <= len(shop) <= 30 else ''


def _get_repeat_analysis(engine, selected_months):
    """分析重复案件：同一街道+同一店铺+同一小类在同一个月内出现多次
    返回: (total, persist, resolved, new, top15_data)
    """
    if not engine or len(selected_months) < 2:
        return 0, 0, 0, 0, []

    m1, m2 = selected_months[0], selected_months[1]
    month_filter = _build_month_filter(selected_months)

    try:
        with engine.connect() as conn:
            df = pd.read_sql(sa_text(f"""
                SELECT street, description, small_category, upload_batch
                FROM case_data {month_filter}
                AND description IS NOT NULL AND description != ''
            """), conn)

        if df.empty:
            return 0, 0, 0, 0, []

        # 提取店铺名称
        df['shop'] = df['description'].apply(_extract_shop_from_desc)
        df_valid = df[df['shop'] != ''].copy()
        df_valid['case_key'] = df_valid['street'] + '_' + df_valid['shop'] + '_' + df_valid['small_category']

        # 获取每月重复的案件组
        def get_repeat_keys(data, month):
            month_data = data[data['upload_batch'] == month]
            key_counts = month_data['case_key'].value_counts()
            return set(key_counts[key_counts >= 2].index)

        may_repeat = get_repeat_keys(df_valid, m1)
        jun_repeat = get_repeat_keys(df_valid, m2)

        persist = may_repeat & jun_repeat
        resolved = may_repeat - jun_repeat
        new = jun_repeat - may_repeat
        total = len(persist) + len(resolved) + len(new)

        # 获取持续存在的Top15
        persist_data = df_valid[df_valid['case_key'].isin(persist)]
        if not persist_data.empty:
            top15 = persist_data.groupby(['street', 'shop', 'small_category']).agg(
                may_count=('upload_batch', lambda x: (x == m1).sum()),
                jun_count=('upload_batch', lambda x: (x == m2).sum())
            ).sort_values('may_count', ascending=False).head(15)
            top15_data = []
            for (street, shop, small_cat), row in top15.iterrows():
                top15_data.append({
                    'shop': f"{street}_{shop}",
                    'may_count': int(row['may_count']),
                    'jun_count': int(row['jun_count'])
                })
        else:
            top15_data = []

        return total, len(persist), len(resolved), len(new), top15_data

    except Exception as e:
        print(f"重复案件分析失败: {e}")
        import traceback
        traceback.print_exc()
        return 0, 0, 0, 0, []


def _get_single_month_repeat_count(engine, selected_months):
    """统计单月内重复案件组数：同一街道+同一店铺+同一小类在同一个月内出现多次
    返回: (repeat_count, top15_data)
    使用与重复次数分布柱状图相同的case_key定义
    """
    if not engine or not selected_months:
        return 0, []
    
    try:
        month_filter = _build_month_filter(selected_months)
        with engine.connect() as conn:
            # 使用与重复次数分布柱状图相同的case_key定义
            df = pd.read_sql(sa_text(f"""
                SELECT CONCAT(LEFT(description,30),'-',small_category) AS case_key,
                       street, description, small_category
                FROM case_data {month_filter}
                AND description IS NOT NULL AND description != ''
            """), conn)
        
        if df.empty:
            return 0, []
        
        # 统计重复组数
        key_counts = df['case_key'].value_counts()
        repeat_groups = key_counts[key_counts >= 2]
        repeat_count = len(repeat_groups)
        
        # 获取Top15重复案件详情
        top15_data = []
        if repeat_count > 0:
            repeat_keys = repeat_groups.index.tolist()
            repeat_data = df[df['case_key'].isin(repeat_keys)]
            top15 = repeat_data.groupby(['case_key', 'street', 'small_category']).agg(
                count=('case_key', 'size')
            ).sort_values('count', ascending=False).head(15)
            
            for (case_key, street, small_cat), row in top15.iterrows():
                # 从case_key中提取店铺名称（前30个字符）
                shop_desc = case_key.split('-')[0] if '-' in case_key else case_key
                top15_data.append({
                    'shop': f"{street}_{shop_desc}",
                    'count': int(row['count'])
                })
        
        return repeat_count, top15_data
        
    except Exception as e:
        print(f"单月重复案件统计失败: {e}")
        return 0, []


def _get_repeat_count(engine, selected_months):
    """Get repeat case group count filtered by selected months"""
    total, _, _, _, _ = _get_repeat_analysis(engine, selected_months)
    return str(total) if total > 0 else "0"

# ============================================================
# Chart generation (matplotlib)
# ============================================================

def _generate_composite_chart(chart_configs, engine, ncols=None):
    """Generate composite chart image (multiple sub-charts in one figure)"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    num_charts = len(chart_configs)
    if num_charts == 0:
        return None
    ncols = ncols or num_charts
    nrows = (num_charts + ncols - 1) // ncols

    fig, axes = plt.subplots(nrows, ncols, figsize=(7 * ncols, 5 * nrows), facecolor='white')
    if num_charts == 1:
        axes = [axes]
    elif hasattr(axes, 'flatten'):
        axes = list(axes.flatten())
    else:
        axes = [axes]

    colors = ['#4472C4', '#ED7D31', '#A5A5A5', '#FFC000', '#5B9BD5', '#70AD47']

    for i, config in enumerate(chart_configs[:num_charts]):
        ax = axes[i]
        ax.set_facecolor('white')
        try:
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(config['query']), conn)
            if df.empty:
                ax.text(0.5, 0.5, '无数据', ha='center', va='center', fontsize=12)
                ax.set_title(config['title'], fontsize=11, fontweight='bold')
                continue
            x_col = df.columns[0]
            y_col = df.columns[1] if len(df.columns) > 1 else None
            chart_type = config.get('chart_type', 'bar')
            x_data = [str(v) for v in df[x_col]]
            y_data = []
            for v in (df[y_col] if y_col else [1]*len(df)):
                if isinstance(v, (list, tuple)): y_data.append(0)
                elif isinstance(v, str):
                    try: y_data.append(float(v))
                    except: y_data.append(0)
                else: y_data.append(float(v) if v is not None else 0)

            if chart_type == 'bar':
                bars = ax.bar(range(len(x_data)), y_data, color=colors[0], width=0.6)
                ax.set_xticks(range(len(x_data)))
                ax.set_xticklabels(x_data, rotation=30, ha='right', fontsize=9)
                for bar, val in zip(bars, y_data):
                    ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,.0f}', ha='center', va='bottom', fontsize=9)
            elif chart_type == 'horizontal_bar':
                y_pos = range(len(x_data))
                bars = ax.barh(y_pos, y_data, color=colors[0], height=0.6)
                ax.set_yticks(y_pos)
                ax.set_yticklabels(x_data, fontsize=9)
                ax.invert_yaxis()
                for bar, val in zip(bars, y_data):
                    ax.text(bar.get_width(), bar.get_y()+bar.get_height()/2., f' {val:,.0f}', va='center', fontsize=9)
            elif chart_type == 'pie':
                labels = df[x_col].astype(str)
                sizes = df[y_col] if y_col else [1]*len(df)
                total = sizes.sum() if sizes.sum() > 0 else 1
                sorted_indices = sorted(range(len(sizes)), key=lambda k: sizes.iloc[k] if hasattr(sizes, 'iloc') else sizes[k], reverse=True)
                top3_set = set(sorted_indices[:3])
                display_labels = [l if i in top3_set else "" for i, l in enumerate(labels)]
                wedges, texts, autotexts = ax.pie(
                    sizes, labels=display_labels,
                    autopct=lambda pct: f"{pct:.1f}%" if pct >= 3 else "",
                    colors=colors[:len(labels)], startangle=90,
                    textprops={'fontsize': 9}, pctdistance=0.75
                )
                for idx_in_top3, global_idx in enumerate(sorted_indices[:3]):
                    texts[global_idx].set_fontsize(10)
                    texts[global_idx].set_fontweight('bold')
            elif chart_type == 'line':
                ax.plot(range(len(x_data)), y_data, color=colors[0], marker='o', linewidth=2, markersize=6)
                ax.fill_between(range(len(x_data)), y_data, alpha=0.15, color=colors[0])
                ax.set_xticks(range(len(x_data)))
                ax.set_xticklabels([str(i+1) for i in range(len(x_data))], fontsize=9)
            elif chart_type in ('compare_bar', 'compare_kpi'):
                # Simple comparison bar chart with 2 bars (2 months)
                bars = ax.bar(range(len(x_data)), y_data, color=[colors[0], colors[1]][:len(x_data)], width=0.5)
                ax.set_xticks(range(len(x_data)))
                # Format month labels (202605 -> 5月)
                month_labels = []
                for m in x_data:
                    if len(m) >= 6:
                        month_labels.append(f"{int(m[4:6])}月")
                    else:
                        month_labels.append(m)
                ax.set_xticklabels(month_labels, fontsize=10)
                for bar, val in zip(bars, y_data):
                    label = f'{val:,.1f}%' if '率' in config['title'] or 'kpi' in chart_type else f'{val:,.0f}'
                    ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), label, ha='center', va='bottom', fontsize=10, fontweight='bold')
            elif chart_type == 'compare_grouped':
                # Grouped bar chart for comparing multiple metrics across 2 months
                if len(df.columns) >= 3:
                    metrics = df.columns[1:]  # e.g., ['延期案件', '返工案件']
                    x_pos = range(len(metrics))
                    bar_width = 0.35
                    for idx, month in enumerate(x_data[:2]):
                        offset = (idx - 0.5) * bar_width
                        values = [float(df.iloc[i][metrics[j]]) if idx < len(df) else 0 for j in range(len(metrics))]
                        bars = ax.bar([p + offset for p in x_pos], values[:len(metrics)], bar_width, label=f"{int(month[4:6])}月" if len(month)>=6 else month, color=colors[idx])
                        for bar, val in zip(bars, values[:len(metrics)]):
                            ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,.0f}', ha='center', va='bottom', fontsize=8)
                    ax.set_xticks(list(x_pos))
                    ax.set_xticklabels(metrics, fontsize=9)
                    ax.legend(fontsize=8)
            elif chart_type == 'compare_grouped_pct':
                # Grouped bar chart for percentage metrics
                if len(df.columns) >= 3:
                    metrics = df.columns[1:]  # e.g., ['延期率', '返工率']
                    x_pos = range(len(metrics))
                    bar_width = 0.35
                    for idx, month in enumerate(x_data[:2]):
                        offset = (idx - 0.5) * bar_width
                        values = [float(df.iloc[i][metrics[j]]) if idx < len(df) else 0 for j in range(len(metrics))]
                        bars = ax.bar([p + offset for p in x_pos], values[:len(metrics)], bar_width, label=f"{int(month[4:6])}月" if len(month)>=6 else month, color=colors[idx])
                        for bar, val in zip(bars, values[:len(metrics)]):
                            ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
                    ax.set_xticks(list(x_pos))
                    ax.set_xticklabels(metrics, fontsize=9)
                    ax.legend(fontsize=8)
            elif chart_type == 'compare_category':
                # Category comparison with grouped bars for each category
                if len(df.columns) >= 3:
                    cat_col = df.columns[0]  # e.g., '大类'
                    month_col = df.columns[1]  # '月份'
                    val_col = df.columns[2]  # '案件数'
                    categories = df[cat_col].unique()[:15]  # Limit to top 15
                    months = sorted(df[month_col].unique())[:2]
                    x_pos = range(len(categories))
                    bar_width = 0.35
                    for idx, month in enumerate(months):
                        month_data = df[df[month_col] == month]
                        values = []
                        for cat in categories:
                            cat_val = month_data[month_data[cat_col] == cat][val_col]
                            values.append(float(cat_val.iloc[0]) if len(cat_val) > 0 else 0)
                        offset = (idx - 0.5) * bar_width
                        bars = ax.barh([p + offset for p in x_pos], values, bar_width, label=f"{int(month[4:6])}月" if len(month)>=6 else month, color=colors[idx])
                        for bar, val in zip(bars, values):
                            if val > 0:
                                ax.text(bar.get_width(), bar.get_y()+bar.get_height()/2., f'{val:,.0f}', va='center', fontsize=7)
                    ax.set_yticks(list(x_pos))
                    ax.set_yticklabels(categories, fontsize=8)
                    ax.invert_yaxis()
                    ax.legend(fontsize=8, loc='lower right')
            else:  # completion_rate, duration_dist etc
                # Check if there's a percentage column (3rd column)
                pct_col = df.columns[2] if len(df.columns) > 2 else None
                if pct_col and '百分比' in pct_col:
                    # Use percentage for bar heights and labels
                    pct_data = []
                    for v in df[pct_col]:
                        if isinstance(v, (list, tuple)): pct_data.append(0)
                        elif isinstance(v, str):
                            try: pct_data.append(float(v))
                            except: pct_data.append(0)
                        else: pct_data.append(float(v) if v is not None else 0)
                    bars = ax.bar(range(len(x_data)), pct_data, color=colors[0], width=0.6)
                    ax.set_xticks(range(len(x_data)))
                    ax.set_xticklabels(x_data, fontsize=9)
                    for bar, val in zip(bars, pct_data):
                        ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
                else:
                    bars = ax.bar(range(len(x_data)), y_data, color=colors[0], width=0.6)
                    ax.set_xticks(range(len(x_data)))
                    ax.set_xticklabels(x_data, fontsize=9)
                    for bar, val in zip(bars, y_data):
                        ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,.0f}', ha='center', va='bottom', fontsize=9)

            ax.set_title(config['title'], fontsize=13, fontweight='bold')
        except Exception as e:
            ax.text(0.5, 0.5, '查询失败', ha='center', va='center', fontsize=10)
            ax.set_title(config['title'], fontsize=13, fontweight='bold')
            print(f"  Sub-chart error: {e}")

    for i in range(num_charts, len(axes)):
        fig.delaxes(axes[i])

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=180, facecolor='white', bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ============================================================
# Comparison chart generation (8 chart types for 2-month comparison)
# ============================================================

def _generate_comparison_chart(chart_config, engine, selected_months):
    """Generate a single comparison chart with exact styling from reference"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    chart_type = chart_config.get('chart_type', '')
    title = chart_config.get('title', '')
    color1, color2 = '#4472C4', '#ED7D31'  # Blue for month1, Orange for month1

    def _get_month_label(m):
        return f"{int(m[4:6])}月" if len(m) >= 6 else m

    m1_label = _get_month_label(selected_months[0]) if len(selected_months) >= 1 else "月1"
    m2_label = _get_month_label(selected_months[1]) if len(selected_months) >= 2 else "月2"

    buf = io.BytesIO()

    try:
        if chart_type == 'compare_total':
            # 图1: 案件总量对比 - simple 2-bar chart
            query = chart_config['query']
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(query), conn)
            fig, ax = plt.subplots(figsize=(8, 5), facecolor='white')
            ax.set_facecolor('white')
            if not df.empty:
                months = [_get_month_label(m) for m in df.iloc[:, 0]]
                values = [float(v) for v in df.iloc[:, 1]]
                bars = ax.bar(months, values, color=[color1, color2][:len(months)], width=0.5)
                for bar, val in zip(bars, values):
                    ax.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,.0f}', ha='center', va='bottom', fontsize=12, fontweight='bold')
            ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
            ax.set_ylabel('案件数', fontsize=11)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            plt.tight_layout()

        elif chart_type == 'compare_efficiency':
            # 图2: 平均处置时长与办结率 - two subplots
            queries = chart_config['queries']
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5), facecolor='white')
            # Left: 平均处置时长
            with engine.connect() as conn:
                df1 = pd.read_sql(sa_text(queries[0]), conn)
            ax1.set_facecolor('white')
            if not df1.empty:
                months = [_get_month_label(m) for m in df1.iloc[:, 0]]
                values = [float(v) for v in df1.iloc[:, 1]]
                bars = ax1.bar(months, values, color=[color1, color2][:len(months)], width=0.5)
                for bar, val in zip(bars, values):
                    ax1.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:.1f}h', ha='center', va='bottom', fontsize=11, fontweight='bold')
            ax1.set_title('平均处置时长', fontsize=13, fontweight='bold')
            ax1.set_ylabel('小时', fontsize=10)
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)
            ax1.grid(axis='y', alpha=0.3)
            # Right: 办结率
            with engine.connect() as conn:
                df2 = pd.read_sql(sa_text(queries[1]), conn)
            ax2.set_facecolor('white')
            if not df2.empty:
                months = [_get_month_label(m) for m in df2.iloc[:, 0]]
                values = [float(v) for v in df2.iloc[:, 1]]
                bars = ax2.bar(months, values, color=[color1, color2][:len(months)], width=0.5)
                for bar, val in zip(bars, values):
                    ax2.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:.1f}%', ha='center', va='bottom', fontsize=11, fontweight='bold')
            ax2.set_title('办结率对比', fontsize=13, fontweight='bold')
            ax2.set_ylabel('%', fontsize=10)
            ax2.set_ylim(0, 110)
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            ax2.grid(axis='y', alpha=0.3)
            plt.suptitle(title, fontsize=14, fontweight='bold', y=1.02)
            plt.tight_layout()

        elif chart_type == 'compare_delay_rework':
            # 图3: 延期与返工数量及比率 - two subplots
            queries = chart_config['queries']
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5), facecolor='white')
            # Left: 案件数
            with engine.connect() as conn:
                df1 = pd.read_sql(sa_text(queries[0]), conn)
            ax1.set_facecolor('white')
            if not df1.empty and len(df1) >= 2:
                categories = ['延期案件', '返工案件']
                x_pos = range(len(categories))
                bar_width = 0.35
                for idx in range(min(2, len(df1))):
                    row = df1.iloc[idx]
                    offset = (idx - 0.5) * bar_width
                    vals = [float(row['延期案件']), float(row['返工案件'])]
                    bars = ax1.bar([p + offset for p in x_pos], vals, bar_width, label=_get_month_label(row['月份']), color=[color1, color2][idx])
                    for bar, val in zip(bars, vals):
                        ax1.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,.0f}', ha='center', va='bottom', fontsize=10)
                ax1.set_xticks(list(x_pos))
                ax1.set_xticklabels(categories, fontsize=10)
                ax1.set_ylabel('案件数', fontsize=10)
                ax1.legend(fontsize=9)
            ax1.set_title('延期与返工案件数', fontsize=13, fontweight='bold')
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)
            # Right: 比率
            with engine.connect() as conn:
                df2 = pd.read_sql(sa_text(queries[1]), conn)
            ax2.set_facecolor('white')
            if not df2.empty and len(df2) >= 2:
                categories = ['延期率', '返工率']
                x_pos = range(len(categories))
                bar_width = 0.35
                for idx in range(min(2, len(df2))):
                    row = df2.iloc[idx]
                    offset = (idx - 0.5) * bar_width
                    vals = [float(row['延期率']), float(row['返工率'])]
                    bars = ax2.bar([p + offset for p in x_pos], vals, bar_width, label=_get_month_label(row['月份']), color=[color1, color2][idx])
                    for bar, val in zip(bars, vals):
                        ax2.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:.2f}%', ha='center', va='bottom', fontsize=9)
                ax2.set_xticks(list(x_pos))
                ax2.set_xticklabels(categories, fontsize=10)
                ax2.set_ylabel('%', fontsize=10)
                ax2.legend(fontsize=9)
            ax2.set_title('延期率与返工率', fontsize=13, fontweight='bold')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            plt.suptitle(title, fontsize=14, fontweight='bold', y=1.02)
            plt.tight_layout()

        elif chart_type == 'compare_category_change':
            # 图4: 大类数量对比及变化率
            query = chart_config['query']
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(query), conn)
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6), facecolor='white')
            if not df.empty:
                cat_col, month_col, val_col = df.columns[0], df.columns[1], df.columns[2]
                months = sorted(df[month_col].unique())[:2]

                # 左图：按两个月总量从高到低排序
                cat_total = df.groupby(cat_col)[val_col].sum().sort_values(ascending=False)
                categories_for_left = cat_total.index.tolist()[:12]

                # 右图：计算变化率，按从正到负排序
                m1_data = df[df[month_col] == months[0]]
                m2_data = df[df[month_col] == months[1]]
                cat_change = []
                for cat in df[cat_col].unique():
                    v1 = float(m1_data[m1_data[cat_col] == cat][val_col].iloc[0]) if not m1_data[m1_data[cat_col] == cat].empty else 0
                    v2 = float(m2_data[m2_data[cat_col] == cat][val_col].iloc[0]) if not m2_data[m2_data[cat_col] == cat].empty else 0
                    rate = ((v2 - v1) / v1 * 100) if v1 > 0 else 0
                    cat_change.append((cat, rate))
                cat_change.sort(key=lambda x: x[1], reverse=True)
                categories_for_right = [c[0] for c in cat_change[:12]]

                # Left: 数量对比（从高到低）
                x_pos_left = range(len(categories_for_left))
                bar_width = 0.35
                for idx, month in enumerate(months):
                    month_data = df[df[month_col] == month]
                    values = [float(month_data[month_data[cat_col] == cat][val_col].iloc[0]) if not month_data[month_data[cat_col] == cat].empty else 0 for cat in categories_for_left]
                    offset = (idx - 0.5) * bar_width
                    ax1.barh([p + offset for p in x_pos_left], values, bar_width, label=_get_month_label(month), color=[color1, color2][idx])
                ax1.set_yticks(list(x_pos_left))
                ax1.set_yticklabels(categories_for_left, fontsize=9)
                ax1.invert_yaxis()
                ax1.set_xlabel('案件数', fontsize=10)
                ax1.legend(fontsize=9)
                ax1.set_title('问题大类数量对比', fontsize=13, fontweight='bold')
                ax1.spines['top'].set_visible(False)
                ax1.spines['right'].set_visible(False)

                # Right: 变化率（从正到负）
                x_pos_right = range(len(categories_for_right))
                change_rates = [c[1] for c in cat_change[:12]]
                bar_colors = ['#70AD47' if r >= 0 else '#FF6B6B' for r in change_rates]
                ax2.barh(x_pos_right, change_rates, color=bar_colors, height=0.5)
                for i, val in enumerate(change_rates):
                    offset = 1.5 if val >= 0 else -1.5
                    ha = 'left' if val >= 0 else 'right'
                    ax2.text(val + offset, i, f'{val:.1f}%', va='center', ha=ha, fontsize=8)
                ax2.set_yticks(list(x_pos_right))
                ax2.set_yticklabels(categories_for_right, fontsize=9)
                ax2.invert_yaxis()
                ax2.set_xlabel('变化率 %', fontsize=10)
                ax2.axvline(x=0, color='gray', linewidth=0.8, linestyle='--')
                ax2.set_title(f'大类变化率 ({_get_month_label(months[1])} vs {_get_month_label(months[0])})', fontsize=13, fontweight='bold')
                ax2.spines['top'].set_visible(False)
                ax2.spines['right'].set_visible(False)
            plt.suptitle(title, fontsize=14, fontweight='bold', y=1.02)
            plt.tight_layout()

        elif chart_type == 'compare_small_category':
            # 图5: 小类Top15变化 - horizontal grouped bars with change annotations
            query = chart_config['query']
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(query), conn)
            fig, ax = plt.subplots(figsize=(14, 8), facecolor='white')
            ax.set_facecolor('white')
            if not df.empty:
                cat_col, month_col, val_col = df.columns[0], df.columns[1], df.columns[2]
                # Get top 15 by total count
                total_by_cat = df.groupby(cat_col)[val_col].sum().sort_values(ascending=False).head(15)
                categories = total_by_cat.index.tolist()
                months = sorted(df[month_col].unique())[:2]
                x_pos = range(len(categories))
                bar_width = 0.35
                for idx, month in enumerate(months):
                    month_data = df[df[month_col] == month]
                    values = []
                    for cat in categories:
                        cat_val = month_data[month_data[cat_col] == cat][val_col]
                        values.append(float(cat_val.iloc[0]) if len(cat_val) > 0 else 0)
                    offset = (idx - 0.5) * bar_width
                    bars = ax.barh([p + offset for p in x_pos], values, bar_width, label=_get_month_label(month), color=[color1, color2][idx])
                # Add change annotations
                m1_data = df[df[month_col] == months[0]] if len(months) >= 1 else pd.DataFrame()
                m2_data = df[df[month_col] == months[1]] if len(months) >= 2 else pd.DataFrame()
                for i, cat in enumerate(categories):
                    v1 = float(m1_data[m1_data[cat_col] == cat][val_col].iloc[0]) if not m1_data[m1_data[cat_col] == cat].empty else 0
                    v2 = float(m2_data[m2_data[cat_col] == cat][val_col].iloc[0]) if not m2_data[m2_data[cat_col] == cat].empty else 0
                    change = v2 - v1
                    pct = (change / v1 * 100) if v1 > 0 else 0
                    max_val = max(v1, v2)
                    ax.text(max_val + 50, i, f'{change:+.0f} ({pct:+.0f}%)', va='center', fontsize=8, color='#70AD47' if change >= 0 else '#FF6B6B')
                ax.set_yticks(list(x_pos))
                ax.set_yticklabels(categories, fontsize=9)
                ax.invert_yaxis()
                ax.set_xlabel('案件数', fontsize=11)
                ax.legend(fontsize=10, loc='lower right')
                ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
            plt.tight_layout()

        elif chart_type == 'compare_repeat':
            # 图6: 重复案件状态及顽固案件 - 使用正确的分析逻辑
            total, persist, resolved, new_count, top15_data = _get_repeat_analysis(engine, selected_months)
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6), facecolor='white')

            # Left: 状态分布
            ax1.set_facecolor('white')
            statuses = ['持续存在', '已解决', '新增']
            values = [persist, resolved, new_count]
            status_colors = ['#E74C3C', '#70AD47', '#ED7D31']
            bars = ax1.bar(statuses, values, color=status_colors, width=0.5)
            for bar, val in zip(bars, values):
                if val > 0:
                    ax1.text(bar.get_x()+bar.get_width()/2., bar.get_height(), f'{val:,}', ha='center', va='bottom', fontsize=12, fontweight='bold')
            ax1.set_title('重复案件状态分布', fontsize=13, fontweight='bold')
            ax1.set_ylabel('案件组数', fontsize=10)
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)

            # Right: 顽固店铺Top15 - 堆叠水平柱状图
            ax2.set_facecolor('white')
            if top15_data:
                shops = [d['shop'] for d in top15_data]
                may_vals = [d['may_count'] for d in top15_data]
                jun_vals = [d['jun_count'] for d in top15_data]
                y_pos = range(len(shops))
                left_vals = [0] * len(shops)
                bars1 = ax2.barh(y_pos, may_vals, left=left_vals, height=0.6, label=_get_month_label(selected_months[0]), color=color1)
                left_vals = [l + v for l, v in zip(left_vals, may_vals)]
                bars2 = ax2.barh(y_pos, jun_vals, left=left_vals, height=0.6, label=_get_month_label(selected_months[1]), color=color2)
                ax2.set_yticks(list(y_pos))
                ax2.set_yticklabels(shops, fontsize=8)
                ax2.invert_yaxis()
                ax2.set_xlabel('出现次数', fontsize=10)
                ax2.legend(fontsize=9)
            ax2.set_title('顽固重复店铺 Top15', fontsize=13, fontweight='bold')
            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            plt.suptitle(title, fontsize=14, fontweight='bold', y=1.02)
            plt.tight_layout()

        elif chart_type == 'compare_department':
            query = chart_config['query']
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(query), conn)
            fig, ax = plt.subplots(figsize=(12, 10), facecolor='white')
            ax.set_facecolor('white')
            if not df.empty:
                dept_col, month_col, val_col = df.columns[0], df.columns[1], df.columns[2]
                # 综合两个月的平均时长，从高到低排序
                dept_avg = df.groupby(dept_col)[val_col].mean().sort_values(ascending=False)
                departments = dept_avg.index.tolist()[:25]
                months = sorted(df[month_col].unique())[:2]
                x_pos = range(len(departments))
                bar_width = 0.35
                for idx, month in enumerate(months):
                    month_data = df[df[month_col] == month]
                    values = []
                    for dept in departments:
                        dept_val = month_data[month_data[dept_col] == dept][val_col]
                        values.append(float(dept_val.iloc[0]) if len(dept_val) > 0 else 0)
                    offset = (idx - 0.5) * bar_width
                    bars = ax.barh([p + offset for p in x_pos], values, bar_width, label=_get_month_label(month), color=[color1, color2][idx])
                ax.set_yticks(list(x_pos))
                ax.set_yticklabels(departments, fontsize=9)
                ax.invert_yaxis()  # 时长最长的在最上面
                ax.set_xlabel('平均时长（小时）', fontsize=11)
                ax.legend(fontsize=10, loc='lower right')
                ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                ax.grid(axis='x', alpha=0.3)
            plt.tight_layout()

        elif chart_type == 'compare_source_change':
            # 图8: 问题来源渠道对比 - two subplots, 按变化率排序
            query = chart_config['query']
            with engine.connect() as conn:
                df = pd.read_sql(sa_text(query), conn)
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6), facecolor='white')
            if not df.empty:
                src_col, month_col, val_col = df.columns[0], df.columns[1], df.columns[2]
                months = sorted(df[month_col].unique())[:2]

                # 按两个月总量从高到低排序（用于左图）
                src_total = df.groupby(src_col)[val_col].sum().sort_values(ascending=False)
                sources = src_total.index.tolist()[:10]

                # 计算变化率（用于右图）
                m1_data = df[df[month_col] == months[0]] if len(months) >= 1 else pd.DataFrame()
                m2_data = df[df[month_col] == months[1]] if len(months) >= 2 else pd.DataFrame()
                src_change = []
                for src in sources:
                    v1 = float(m1_data[m1_data[src_col] == src][val_col].iloc[0]) if not m1_data[m1_data[src_col] == src].empty else 0
                    v2 = float(m2_data[m2_data[src_col] == src][val_col].iloc[0]) if not m2_data[m2_data[src_col] == src].empty else 0
                    rate = ((v2 - v1) / v1 * 100) if v1 > 0 else 0
                    src_change.append((src, v1, v2, rate))

                # Left: 数量对比（按总量从高到低）
                x_pos = range(len(sources))
                bar_width = 0.35
                for idx, month in enumerate(months):
                    month_data = df[df[month_col] == month]
                    values = []
                    for src in sources:
                        src_val = month_data[month_data[src_col] == src][val_col]
                        values.append(float(src_val.iloc[0]) if len(src_val) > 0 else 0)
                    offset = (idx - 0.5) * bar_width
                    bars = ax1.barh([p + offset for p in x_pos], values, bar_width, label=_get_month_label(month), color=[color1, color2][idx])
                ax1.set_yticks(list(x_pos))
                ax1.set_yticklabels(sources, fontsize=9)
                ax1.invert_yaxis()
                ax1.set_xlabel('案件数', fontsize=10)
                ax1.legend(fontsize=9)
                ax1.set_title('问题来源对比', fontsize=13, fontweight='bold')
                ax1.spines['top'].set_visible(False)
                ax1.spines['right'].set_visible(False)

                # Right: 变化率（按从正到负排序）
                src_change.sort(key=lambda x: x[3], reverse=True)
                sources_right = [s[0] for s in src_change[:10]]
                change_rates = [s[3] for s in src_change[:10]]
                x_pos_right = range(len(sources_right))
                bar_colors = ['#70AD47' if r >= 0 else '#FF6B6B' for r in change_rates]
                ax2.barh(x_pos_right, change_rates, color=bar_colors, height=0.5)
                for i, val in enumerate(change_rates):
                    offset = 1.5 if val >= 0 else -1.5
                    ha = 'left' if val >= 0 else 'right'
                    ax2.text(val + offset, i, f'{val:.1f}%', va='center', ha=ha, fontsize=8)
                ax2.set_yticks(list(x_pos_right))
                ax2.set_yticklabels(sources_right, fontsize=9)
                ax2.invert_yaxis()
                ax2.set_xlabel('变化率 %', fontsize=10)
                ax2.axvline(x=0, color='gray', linewidth=0.8, linestyle='--')
                ax2.set_title(f'问题来源变化率 ({_get_month_label(months[1])} vs {_get_month_label(months[0])})', fontsize=13, fontweight='bold')
                ax2.spines['top'].set_visible(False)
                ax2.spines['right'].set_visible(False)
            plt.suptitle(title, fontsize=14, fontweight='bold', y=1.02)
            plt.tight_layout()

        else:
            # Fallback
            fig, ax = plt.subplots(figsize=(8, 5), facecolor='white')
            ax.text(0.5, 0.5, '图表类型未实现', ha='center', va='center', fontsize=12)
            ax.set_title(title, fontsize=13, fontweight='bold')

        fig.savefig(buf, format='png', dpi=180, facecolor='white', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        print(f"Comparison chart error: {e}")
        plt.close('all')
        return None

# ============================================================
# Placeholder replacement
# ============================================================

def _fill_placeholders(doc, year, month_num, summary_data, engine=None, selected_months=None, report_type='single'):
    """Replace placeholders in document - handles cross-run placeholders"""
    import re

    def _do_replacements(text):
        new_text = text
        # Comparison report placeholders
        if report_type == 'compare' and len(selected_months) >= 2:
            # 确保月份按时间顺序排列：earlier_month 在前，later_month 在后
            sorted_months = sorted(selected_months)
            earlier_month = sorted_months[0]  # 前月（如202605）
            later_month = sorted_months[1]    # 后月（如202606）
            earlier_num = str(int(earlier_month[4:6])) if len(earlier_month) >= 6 else "?"
            later_num = str(int(later_month[4:6])) if len(later_month) >= 6 else "?"

            # 获取两个月的分开数据（前月数据在前，后月数据在后）
            earlier_filter = _build_month_filter([earlier_month])
            later_filter = _build_month_filter([later_month])
            try:
                with engine.connect() as conn:
                    earlier_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),1) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {earlier_filter}"
                    later_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),1) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {later_filter}"
                    earlier_df = pd.read_sql(sa_text(earlier_sql), conn)
                    later_df = pd.read_sql(sa_text(later_sql), conn)
                    earlier_data = earlier_df.iloc[0].to_dict() if not earlier_df.empty else {}
                    later_data = later_df.iloc[0].to_dict() if not later_df.empty else {}
            except Exception as e:
                print(f"获取对比数据失败: {e}")
                earlier_data, later_data = {}, {}

            # 前月数据
            t1 = int(earlier_data.get('案件总量', 0))
            avg1 = earlier_data.get('平均时长', 0)
            cr1 = earlier_data.get('办结率', 0)
            d1 = int(earlier_data.get('延期案件', 0))
            r1 = int(earlier_data.get('返工案件', 0))
            # 后月数据
            t2 = int(later_data.get('案件总量', 0))
            avg2 = later_data.get('平均时长', 0)
            cr2 = later_data.get('办结率', 0)
            d2 = int(later_data.get('延期案件', 0))
            r2 = int(later_data.get('返工案件', 0))

            # 获取重复案件数据
            total_r, persist_r, resolved_r, new_r, _ = _get_repeat_analysis(engine, selected_months)

            # 计算变化百分比（后月 vs 前月）
            total_pct = ((t2 - t1) / t1 * 100) if t1 > 0 else 0

            # 替换标题模式
            new_text = new_text.replace('（）月 vs （）月', f'{earlier_num}月 vs {later_num}月')
            new_text = new_text.replace('（）年（）月-（）月', f'{earlier_month[:4]}年{earlier_num}月-{later_num}月')
            new_text = new_text.replace('（）月较（）月', f'{later_num}月较{earlier_num}月')

            # 替换月份数字（前月在前，后月在后）
            new_text = new_text.replace('（）月', f'{earlier_num}月', 1)  # 第一个（）月=前月
            new_text = new_text.replace('（）月', f'{later_num}月', 1)   # 第二个（）月=后月

            # 替换具体数据
            new_text = new_text.replace('减少（）%', f'减少{abs(total_pct):.1f}%')
            new_text = new_text.replace('增加（）%', f'增加{abs(total_pct):.1f}%')
            new_text = new_text.replace('从（）延长至（）', f'从{avg1}h延长至{avg2}h')
            new_text = new_text.replace('从（）h延长至（）h', f'从{avg1}h延长至{avg2}h')

            # 替换重复案件数据
            import re
            new_text = re.sub(r'共发现\s+组', f'共发现 {total_r:,} 组', new_text)
            new_text = new_text.replace('持续存在 （）组', f'持续存在 {persist_r:,} 组')
            new_text = new_text.replace('持续存在（）组', f'持续存在 {persist_r:,} 组')
            new_text = new_text.replace('已解决 （）组', f'已解决 {resolved_r:,} 组')
            new_text = new_text.replace('已解决（）组', f'已解决 {resolved_r:,} 组')
            new_text = new_text.replace('新增 （）组', f'新增 {new_r:,} 组')
            new_text = new_text.replace('新增（）组', f'新增 {new_r:,} 组')

            # 替换总结中的数据（前月→后月）
            new_text = new_text.replace(f'{earlier_num}月（）件', f'{earlier_num}月{t1:,}件')
            new_text = new_text.replace(f'{later_num}月（）件', f'{later_num}月{t2:,}件')
            new_text = new_text.replace('（）件→', f'{t1:,}件→')
            new_text = new_text.replace('→（）件', f'→{t2:,}件')
            new_text = new_text.replace('→（）h', f'→{avg2}h')
            new_text = new_text.replace('（）→（）h', f'{avg1}h→{avg2}h')
            new_text = new_text.replace('办结率（）%', f'办结率{cr2}%')
            new_text = new_text.replace('延期（）→（）', f'延期{d1}→{d2}')
            new_text = new_text.replace('返工（）→（）', f'返工{r1}→{r2}')

        # Single month placeholders
        new_text = new_text.replace('（）年（）月', f'{year}年{month_num}月')
        new_text = new_text.replace('数据量：（）件', f'数据量：{summary_data.get("案件总量", "-")}件')

        # 单月报告重复案件数据填充
        if report_type == 'single' and engine and selected_months:
            single_repeat_count, _ = _get_single_month_repeat_count(engine, selected_months)
            # 使用正则表达式匹配"共发现"和"组"之间的任意空格
            import re
            new_text = re.sub(r'共发现\s+组', f'共发现 {single_repeat_count:,} 组', new_text)

        # 通用占位符替换
        new_text = new_text.replace('（）', '')

        return new_text

    def _replace_paragraph_text(para, new_full_text):
        if not para.runs:
            return
        para.runs[0].text = new_full_text
        for run in para.runs[1:]:
            run.text = ''

    for para in doc.paragraphs:
        # Skip title paragraph (already handled by _update_document_title)
        if para.style.name in ('Title', 'Title 1', 'Title 2'):
            continue
        full_text = ''.join(run.text for run in para.runs)
        new_text = _do_replacements(full_text)
        if new_text != full_text:
            _replace_paragraph_text(para, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = ''.join(run.text for run in para.runs)
                    new_text = _do_replacements(full_text)
                    if new_text != full_text:
                        _replace_paragraph_text(para, new_text)

# ============================================================
# Chart insertion (uses python-docx standard API)
# ============================================================

def _insert_image_after_paragraph(doc, paragraph, img_bytes, width_inches=6.8):
    """Insert image after a paragraph using python-docx standard API, centered"""
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    img_stream = io.BytesIO(img_bytes)
    # Use python-docx add_picture to create proper drawing XML
    temp_para = doc.add_paragraph()
    temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = temp_para.add_run()
    run.add_picture(img_stream, width=Inches(width_inches))

    drawing_elements = temp_para._element.findall('.//' + qn('w:drawing'))
    if not drawing_elements:
        temp_para._element.getparent().remove(temp_para._element)
        return

    drawing = drawing_elements[0]
    new_para = OxmlElement('w:p')
    # Set center alignment on the new paragraph
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_para.append(pPr)
    new_run = OxmlElement('w:r')
    new_run.append(copy.deepcopy(drawing))
    new_para.append(new_run)
    paragraph._element.addnext(new_para)
    temp_para._element.getparent().remove(temp_para._element)


# ============================================================
# Section charts: replace description paragraphs with chart images
# ============================================================


def _add_caption_after(doc, reference_element, caption_text):
    """Insert a centered caption paragraph after the reference XML element."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    new_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_para.append(pPr)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '666666')
    rPr.append(color)
    new_run.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = caption_text
    new_run.append(t_elem)
    new_para.append(new_run)
    reference_element.addnext(new_para)


def _fill_section_charts(doc, engine, selected_months, report_type='single'):
    """Generate charts and replace description paragraphs. Two-pass: collect then insert forwards for correct numbering."""
    paragraphs = list(doc.paragraphs)

    # Pass 1 (backwards): collect work items and generate chart images
    work_items = []  # each: (desc_para, [(cfg, img_bytes), ...])
    for i in range(len(paragraphs) - 1, -1, -1):
        para = paragraphs[i]
        if para.style.name == 'Heading 1':
            section_title = para.text.strip()
            chart_configs = _get_chart_configs_with_month(section_title, selected_months, report_type)
            if chart_configs:
                desc_para = None
                for j in range(i + 1, min(i + 3, len(paragraphs))):
                    if paragraphs[j].text.strip():
                        desc_para = paragraphs[j]
                        break
                if desc_para is not None:
                    is_compare = any(c.get('chart_type', '').startswith('compare_') for c in chart_configs)
                    charts = []
                    try:
                        if is_compare:
                            for cfg in chart_configs:
                                img_bytes = _generate_comparison_chart(cfg, engine, selected_months)
                                if img_bytes:
                                    charts.append((cfg, img_bytes))
                        else:
                            num_charts = len(chart_configs)
                            img_bytes = _generate_composite_chart(chart_configs, engine, ncols=num_charts)
                            if img_bytes:
                                charts.append((chart_configs[0], img_bytes))
                    except Exception as e:
                        print(f"  Chart gen error for {section_title}: {e}")
                    if charts:
                        work_items.append((desc_para, charts))

    # Pass 2 (forwards): insert images and captions with correct numbering
    work_items.reverse()
    fig_num = 0
    for desc_para, charts in work_items:
        for cfg, img_bytes in charts:
            fig_num += 1
            try:
                _insert_image_after_paragraph(doc, desc_para, img_bytes)
                img_elem = desc_para._element.getnext()
                if img_elem is not None:
                    from docx.oxml.ns import qn as _qn
                    if img_elem.findall('.//' + _qn('w:drawing')):
                        _add_caption_after(doc, img_elem, f"图{fig_num}: {cfg['title']}")
            except Exception as e:
                print(f"  Chart insert error for '{cfg.get('title')}': {e}")
        try:
            desc_para._element.getparent().remove(desc_para._element)
        except Exception:
            pass
    print(f"  Total charts inserted: {fig_num}")

# ============================================================
# Table data filling
# ============================================================

def _replace_table_data(doc, table_index, new_data, max_rows=20):
    """Replace table data with new rows"""
    if table_index >= len(doc.tables) or not new_data:
        return
    table = doc.tables[table_index]
    cols = list(new_data[0].keys())
    # Clear existing data rows (keep header)
    while len(table.rows) > 1:
        row = table.rows[-1]
        row._element.getparent().remove(row._element)
    # Add new data rows
    for row_data in new_data[:max_rows]:
        row = table.add_row()
        for j, col in enumerate(cols):
            if j < len(row.cells):
                val = row_data.get(col, "")
                if isinstance(val, float):
                    row.cells[j].text = f"{val:,.2f}"
                else:
                    row.cells[j].text = str(val) if val else ""


def _fill_tables(doc, engine, selected_months, summary_data=None, report_type='single'):
    """Fill table data in template"""
    for table_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]

        if '案件总量' in header and '办结率' in header:
            # Summary table - fill row 2 with actual data
            if len(table.rows) >= 2:
                data_row = table.rows[1]
                total = str(summary_data.get("案件总量", "-")) if summary_data else "-"
                avg_duration = str(summary_data.get("平均处置时长", "-")) if summary_data else "-"
                closure_rate = str(summary_data.get("办结率", "-")) if summary_data else "-"
                delayed = str(summary_data.get("延期案件", "-")) if summary_data else "-"
                rework = str(summary_data.get("返工案件", "-")) if summary_data else "-"
                values = [total, avg_duration, closure_rate, delayed, rework]
                for j, val in enumerate(values):
                    if j < len(data_row.cells):
                        data_row.cells[j].text = str(val)

        elif '排名' in header and '地点' in header:
            # 重复案件 Top12 table - 使用正确的分析逻辑
            try:
                if report_type == 'compare' and len(selected_months) >= 2:
                    _, _, _, _, top15_data = _get_repeat_analysis(engine, selected_months)
                else:
                    _, top15_data = _get_single_month_repeat_count(engine, selected_months)
                
                if top15_data:
                    table_data = []
                    for idx, item in enumerate(top15_data[:12]):
                        # 处理不同函数返回的数据格式
                        if 'may_count' in item:
                            # 对比报告格式
                            count = item['may_count'] + item['jun_count']
                        else:
                            # 单月报告格式
                            count = item['count']
                        table_data.append({
                            '排名': idx + 1,
                            '地点': item['shop'],
                            '次数': count
                        })
                    _replace_table_data(doc, table_idx, table_data)
            except Exception as e:
                print(f"  Repeat cases table error: {e}")

        elif '指标' in header and '变化' in header and len(selected_months) >= 2:
            # 确保月份按时间顺序排列
            sorted_months = sorted(selected_months)
            earlier_month = sorted_months[0]  # 前月
            later_month = sorted_months[1]    # 后月
            earlier_filter = _build_month_filter([earlier_month])
            later_filter = _build_month_filter([later_month])
            try:
                with engine.connect() as conn:
                    earlier_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),2) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, ROUND(SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),3) AS 延期率, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件, ROUND(SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),3) AS 返工率 FROM case_data {earlier_filter}"
                    later_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),2) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, ROUND(SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),3) AS 延期率, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件, ROUND(SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END)*100.0/COUNT(*),3) AS 返工率 FROM case_data {later_filter}"
                    earlier_df = pd.read_sql(sa_text(earlier_sql), conn)
                    later_df = pd.read_sql(sa_text(later_sql), conn)
                    if not earlier_df.empty and not later_df.empty:
                        earlier = earlier_df.iloc[0]  # 前月数据
                        later = later_df.iloc[0]      # 后月数据
                        metrics = [
                            ('延期案件', '延期案件', False),
                            ('延期率', '延期率', True),
                            ('返工案件', '返工案件', False),
                            ('返工率', '返工率', True),
                        ]
                        for ri, (metric_name, key, is_pct) in enumerate(metrics):
                            if ri + 1 < len(table.rows):
                                row = table.rows[ri + 1]
                                v1 = earlier[key]  # 前月数据
                                v2 = later[key]    # 后月数据
                                if len(row.cells) >= 4:
                                    if is_pct:
                                        row.cells[1].text = f'{float(v1):.2f}%' if float(v1) < 1 else f'{float(v1):.1f}%'
                                        row.cells[2].text = f'{float(v2):.2f}%' if float(v2) < 1 else f'{float(v2):.1f}%'
                                        change = float(v2) - float(v1)
                                        row.cells[3].text = f'{change:+.2f}%' if abs(change) < 1 else f'{change:+.1f}%'
                                    else:
                                        v1_int, v2_int = int(v1), int(v2)
                                        row.cells[1].text = str(v1_int)
                                        row.cells[2].text = str(v2_int)
                                        row.cells[3].text = str(v2_int - v1_int)
                        # 表头：前月在左，后月在右
                        earlier_num = str(int(earlier_month[4:6])) if len(earlier_month) >= 6 else "月1"
                        later_num = str(int(later_month[4:6])) if len(later_month) >= 6 else "月2"
                        header_row = table.rows[0]
                        if len(header_row.cells) >= 3:
                            header_row.cells[1].text = f'{earlier_num}月'
                            header_row.cells[2].text = f'{later_num}月'
            except Exception as e:
                print(f"  Comparison table error: {e}")

        elif ('顽固' in ' '.join(header) or 'Top8' in ' '.join(header)) and len(selected_months) >= 2:
            m1_label = f"{int(selected_months[0][4:6])}月" if len(selected_months[0]) >= 6 else "月1"
            m2_label = f"{int(selected_months[1][4:6])}月" if len(selected_months[1]) >= 6 else "月2"
            try:
                # 使用正确的重复案件分析逻辑
                _, persist_r, _, _, top15_data = _get_repeat_analysis(engine, selected_months)
                if top15_data:
                    pivot_data = []
                    for item in top15_data[:8]:  # Top8
                        shop = item['shop']
                        # 提取大类（从case_key中无法直接获取，需要额外查询）
                        pivot_data.append({
                            '店铺': shop,
                            '大类': '街面秩序',  # 默认大类
                            m1_label: item['may_count'],
                            m2_label: item['jun_count']
                        })
                    if pivot_data:
                        header_row = table.rows[0]
                        if len(header_row.cells) >= 4:
                            header_row.cells[0].text = '顽固案件 Top8'
                            header_row.cells[1].text = '大类'
                            header_row.cells[2].text = m1_label
                            header_row.cells[3].text = m2_label
                        while len(table.rows) > 1:
                            table.rows[-1]._element.getparent().remove(table.rows[-1]._element)
                        for rd in pivot_data:
                            row = table.add_row()
                            row.cells[0].text = rd['店铺']
                            row.cells[1].text = rd['大类']
                            row.cells[2].text = str(rd[m1_label])
                            row.cells[3].text = str(rd[m2_label])
                        print(f"  顽固案件 Top8 table filled ({len(pivot_data)} rows)")
            except Exception as e:
                print(f"  顽固案件 table error: {e}")

# ============================================================
# Summary section generation (LLM-powered)
# ============================================================

def _fill_summary_section(doc, engine, selected_months, summary_data, report_type='single'):
    """Generate summary & recommendations section using LLM"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    key_data = {}
    month_filter = _build_month_filter(selected_months)

    if report_type == 'compare' and len(selected_months) >= 2:
        # 确保月份按时间顺序排列
        sorted_months = sorted(selected_months)
        earlier_month = sorted_months[0]  # 前月
        later_month = sorted_months[1]    # 后月
        earlier_filter = _build_month_filter([earlier_month])
        later_filter = _build_month_filter([later_month])
        earlier_label = f"{int(earlier_month[4:6])}月" if len(earlier_month) >= 6 else "月1"
        later_label = f"{int(later_month[4:6])}月" if len(later_month) >= 6 else "月2"
        try:
            with engine.connect() as conn:
                earlier_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),1) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {earlier_filter}"
                later_sql = f"SELECT COUNT(*) AS 案件总量, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, ROUND(COUNT(CASE WHEN close_time IS NOT NULL THEN 1 END)*100.0/COUNT(*),1) AS 办结率, SUM(CASE WHEN is_delayed=1 THEN 1 ELSE 0 END) AS 延期案件, SUM(CASE WHEN is_rework=1 THEN 1 ELSE 0 END) AS 返工案件 FROM case_data {later_filter}"
                earlier_df = pd.read_sql(sa_text(earlier_sql), conn)
                later_df = pd.read_sql(sa_text(later_sql), conn)
                if not earlier_df.empty:
                    key_data['m1'] = earlier_df.iloc[0].to_dict()  # 前月数据
                if not later_df.empty:
                    key_data['m2'] = later_df.iloc[0].to_dict()  # 后月数据

                cat_sql = f"SELECT big_category AS 大类, upload_batch AS 月份, COUNT(*) AS 案件数 FROM case_data {month_filter} GROUP BY big_category, upload_batch ORDER BY 案件数 DESC"
                cat_df = pd.read_sql(sa_text(cat_sql), conn)
                if not cat_df.empty:
                    key_data['category_change'] = cat_df.to_dict('records')

                dept_sql = f"SELECT department AS 部门, upload_batch AS 月份, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) AS 平均时长, COUNT(*) AS 案件数 FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY department, upload_batch ORDER BY 案件数 DESC LIMIT 10"
                dept_df = pd.read_sql(sa_text(dept_sql), conn)
                if not dept_df.empty:
                    key_data['dept_change'] = dept_df.to_dict('records')

                # 使用正确的重复案件分析逻辑
                total_r, persist_r, resolved_r, new_r, _ = _get_repeat_analysis(engine, selected_months)
                key_data['repeat_total'] = total_r
                key_data['repeat_persist'] = persist_r
                key_data['repeat_resolved'] = resolved_r
                key_data['repeat_new'] = new_r
        except Exception as e:
            print(f"  Comparison data error: {e}")

        m1 = key_data.get('m1', {})  # 前月数据
        m2 = key_data.get('m2', {})  # 后月数据

        cat_lines = ""
        if 'category_change' in key_data:
            cats = {}
            for r in key_data['category_change']:
                c = r.get('大类', '?')
                m = r.get('月份', '')
                v = r.get('案件数', 0)
                if c not in cats:
                    cats[c] = {}
                cats[c][m] = v
            for c, mv in list(cats.items())[:5]:
                v1 = mv.get(earlier_month, 0)  # 前月数据
                v2 = mv.get(later_month, 0)    # 后月数据
                pct = ((v2 - v1) / v1 * 100) if v1 > 0 else 0
                cat_lines += f"  - {c}: {earlier_label}{v1}件 → {later_label}{v2}件 ({pct:+.1f}%)\n"

        t1 = m1.get('案件总量', 0) or 0  # 前月
        t2 = m2.get('案件总量', 0) or 0  # 后月
        avg1 = m1.get('平均时长', 0) or 0  # 前月
        avg2 = m2.get('平均时长', 0) or 0  # 后月
        cr1 = m1.get('办结率', 0) or 0
        cr2 = m2.get('办结率', 0) or 0
        d1 = m1.get('延期案件', 0) or 0
        d2 = m2.get('延期案件', 0) or 0
        r1 = m1.get('返工案件', 0) or 0
        r2 = m2.get('返工案件', 0) or 0

        try:
            t1, t2 = int(t1), int(t2)
            avg1, avg2 = float(avg1), float(avg2)
            cr1, cr2 = float(cr1), float(cr2)
            d1, d2 = int(d1), int(d2)
            r1, r2 = int(r1), int(r2)
        except (ValueError, TypeError):
            pass

        total_pct = ((t2 - t1) / t1 * 100) if isinstance(t1, (int, float)) and t1 > 0 else 0

        # 使用新的重复案件数据
        persist_r = key_data.get('repeat_persist', 0)
        resolved_r = key_data.get('repeat_resolved', 0)
        new_r = key_data.get('repeat_new', 0)
        total_r = key_data.get('repeat_total', 0)
        repeat_lines = f"持续存在{persist_r}组、已解决{resolved_r}组、新增{new_r}组" if total_r > 0 else ""

        summary_points = [
            f"1. 案件总量：{later_label}较{earlier_label}{'增加' if total_pct >= 0 else '减少'}{abs(total_pct):.1f}%（{earlier_label}{t1:,}件，{later_label}{t2:,}件）。",
            f"2. 处置效率：平均时长{avg1}h→{avg2}h，办结率{cr2}%。",
            f"3. 延期返工：延期{d1}→{d2}件，返工{r1}→{r2}件。",
        ]
        if repeat_lines:
            summary_points.append(f"4. 重复案件：{repeat_lines}。")

        summary_text = "\n".join(summary_points)

        prompt = f"""你是城市管理案件数据分析专家，基于以下对比数据，生成5条具体建议（每条一句话），用于对比分析报告的"建议"部分。

## 数据概况
- {earlier_label}：案件{t1}件，平均时长{avg1}h，办结率{cr1}%，延期{d1}件，返工{r1}件
- {later_label}：案件{t2}件，平均时长{avg2}h，办结率{cr2}%，延期{d2}件，返工{r2}件
- 重复案件：{repeat_lines}

## 大类变化
{cat_lines}

## 输出格式
直接输出5条建议，每条一行，以"-"开头，不要编号，不要标题行。"""

        llm_result = _call_llm([{"role": "user", "content": [{"type": "text", "text": prompt}]}], timeout=60)

        suggestions = []
        if llm_result:
            for line in llm_result.split('\n'):
                stripped = line.strip().lstrip('- ').lstrip('• ').strip()
                if stripped and len(stripped) > 5:
                    suggestions.append(stripped)
        if not suggestions:
            suggestions = [
                f"重点关注{repeat_lines.split('、')[0] if repeat_lines else '持续存在'}的店铺，加强巡查频次。",
                "分析延期案件根因，优化处置流程，缩短响应时间。",
                "关注案件数量增长较快的问题类别，开展专项整治。",
                "优化处置效率偏低的部门资源配置和培训。",
            ]

        sections = {'整体评估': summary_text, '重点工作': '', '下月关注': ''}

        paragraphs = list(doc.paragraphs)
        heading2_map = {}
        for para in paragraphs:
            if para.style.name == 'Heading 2':
                text = para.text.strip()
                if '整体评估' in text or '对比总述' in text or '总述' in text or '数据概览' in text:
                    heading2_map['整体评估'] = para
                elif '重点' in text:
                    heading2_map['重点工作'] = para
                elif '下阶段' in text or '下月' in text or '建议' in text:
                    heading2_map['下月关注'] = para

        def _make_text_para(text_content, font_size='22'):
            new_para = OxmlElement('w:p')
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), font_size)
            rPr.append(sz)
            new_run.append(rPr)
            t_elem = OxmlElement('w:t')
            t_elem.set(qn('xml:space'), 'preserve')
            t_elem.text = text_content
            new_run.append(t_elem)
            new_para.append(new_run)
            return new_para

        for key, heading_para in heading2_map.items():
            if key == '整体评估':
                lines = [l.strip() for l in summary_text.split('\n') if l.strip()]
                last_elem = heading_para._element
                for line in lines:
                    p = _make_text_para(line)
                    last_elem.addnext(p)
                    last_elem = p
                print(f"  Summary inserted: {key} ({len(lines)} lines)")
            elif key == '下月关注':
                last_elem = heading_para._element
                for s in suggestions:
                    p = _make_text_para(f"• {s}")
                    last_elem.addnext(p)
                    last_elem = p
                print(f"  Summary inserted: {key} ({len(suggestions)} suggestions)")
        return

    # --- Single month summary (original logic) ---
    month_filter_and = month_filter.replace('WHERE', 'AND', 1) if month_filter.startswith('WHERE') else month_filter
    try:
        with engine.connect() as conn:
            df = pd.read_sql(sa_text(f"SELECT big_category, COUNT(*) as cnt, ROUND(COUNT(*)/(SELECT COUNT(*) FROM case_data {month_filter})*100,1) as pct FROM case_data {month_filter} GROUP BY big_category ORDER BY cnt DESC LIMIT 5"), conn)
            if not df.empty:
                key_data['top_categories'] = df.to_dict('records')
    except Exception as e:
        print(f"  Category data error: {e}")

    try:
        with engine.connect() as conn:
            df = pd.read_sql(sa_text(f"SELECT department, COUNT(*) as cnt, ROUND(AVG(TIMESTAMPDIFF(MINUTE, report_time, close_time))/60.0,1) as avg_hours FROM case_data {month_filter} AND close_time IS NOT NULL GROUP BY department ORDER BY avg_hours DESC LIMIT 5"), conn)
            if not df.empty:
                key_data['slow_departments'] = df.to_dict('records')
    except Exception as e:
        print(f"  Department data error: {e}")

    total = summary_data.get('案件总量', '-')
    avg_time = summary_data.get('平均处置时长', '-')
    closure = summary_data.get('办结率', '-')
    delayed = summary_data.get('延期案件', '-')
    rework = summary_data.get('返工案件', '-')

    top_cats_str = ""
    if 'top_categories' in key_data:
        items = [f"{r.get('big_category','?')}{r.get('cnt',0)}件({r.get('pct',0)}%)" for r in key_data['top_categories']]
        top_cats_str = "、".join(items)

    slow_depts_str = ""
    if 'slow_departments' in key_data:
        items = [f"{r.get('department','?')}平均{r.get('avg_hours',0)}h" for r in key_data['slow_departments'][:3]]
        slow_depts_str = "、".join(items)

    prompt = f"""你是城市管理案件数据分析专家，正在撰写月度数据分析报告的总结与建议部分。

## 本月数据概况
- 案件总量：{total}件
- 平均处置时长：{avg_time}
- 办结率：{closure}
- 延期案件：{delayed}件
- 返工案件：{rework}件
- 问题大类分布：{top_cats_str}
- 处置较慢部门：{slow_depts_str}

## 输出要求
请生成三个子章节，每个子章节3-5句话，要求：
- 有具体数据支撑，不要空话套话
- 针对性强，指出具体问题和改进方向
- 语言正式，适合政府工作报告

## 输出格式

一、整体评估
（评价本月案件整体情况，突出关键指标变化和亮点）

二、重点工作
（分析需要重点关注的领域、问题及原因，特别是延期和返工情况）

三、下月关注
（提出下月工作重点和具体改进建议）"""

    llm_result = _call_llm([{"role": "user", "content": [{"type": "text", "text": prompt}]}], timeout=60)

    if not llm_result:
        llm_result = f"""一、整体评估
本月共处理案件{total}件，办结率{closure}，平均处置时长{avg_time}。{'延期案件较多，需重点关注。' if isinstance(delayed, int) and delayed > 100 else '整体运行平稳。'}

二、重点工作
{'延期案件{0}件，建议分析延期原因并制定催办机制。'.format(delayed) if isinstance(delayed, int) and delayed > 0 else '案件处理及时性较好。'}{'返工案件{0}件，需关注返工率较高的案件类型。'.format(rework) if isinstance(rework, int) and rework > 0 else ''}

三、下月关注
建议持续提升办结率、缩短平均处置时长，重点关注延期和返工案件的根因分析和流程优化。"""

    # Parse LLM result into sections
    sections = {}
    current_key = None
    current_text = []
    for line in llm_result.split('\n'):
        stripped = line.strip()
        if stripped.startswith('一、整体评估'):
            if current_key: sections[current_key] = '\n'.join(current_text).strip()
            current_key = '整体评估'
            current_text = []
        elif stripped.startswith('二、重点工作'):
            if current_key: sections[current_key] = '\n'.join(current_text).strip()
            current_key = '重点工作'
            current_text = []
        elif stripped.startswith('三、下月关注'):
            if current_key: sections[current_key] = '\n'.join(current_text).strip()
            current_key = '下月关注'
            current_text = []
        elif current_key:
            if stripped and not stripped.endswith('：') and not stripped.endswith(':'):
                current_text.append(stripped)
    if current_key:
        sections[current_key] = '\n'.join(current_text).strip()

    # Find Heading 2 subheadings and insert content after them
    paragraphs = list(doc.paragraphs)
    heading2_map = {}
    for para in paragraphs:
        if para.style.name == 'Heading 2':
            text = para.text.strip()
            if '整体评估' in text: heading2_map['整体评估'] = para
            elif '重点工作' in text: heading2_map['重点工作'] = para
            elif '下月关注' in text: heading2_map['下月关注'] = para

    for key, heading_para in heading2_map.items():
        content = sections.get(key, '')
        if not content:
            continue
        new_para = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '22')
        rPr.append(sz)
        new_run.append(rPr)
        new_text_elem = OxmlElement('w:t')
        new_text_elem.set(qn('xml:space'), 'preserve')
        new_text_elem.text = content
        new_run.append(new_text_elem)
        new_para.append(new_run)
        heading_para._element.addnext(new_para)
        print(f"  Summary inserted: {key} ({len(content)} chars)")

# ============================================================
# Report title helpers
# ============================================================

def _build_report_title(template_name, year, month_num, selected_months):
    """Build a proper report title from template name + selected period"""
    base = template_name.replace('模板', '').strip()
    if year and month_num:
        if len(selected_months) and len(selected_months) > 1:
            months_display = '、'.join(
                f"{m[:4]}年{int(m[4:6])}月" for m in selected_months
                if len(m) >= 6
            )
            return f"{months_display}{base}"
        return f"{year}年{month_num}月{base}"
    return base


def _update_document_title(doc, report_title, original_template_name):
    """Update the document's title paragraph to use the actual report title"""
    clean_template = original_template_name.replace('模板', '').strip()
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # For Title style paragraphs, replace the entire content
        if para.style.name in ('Title', 'Title 1', 'Title 2'):
            if para.runs:
                para.runs[0].text = report_title
                for r in para.runs[1:]:
                    r.text = ''
            return
        # Check first non-empty paragraph if it contains the template name
        if clean_template in text:
            for run in para.runs:
                if clean_template in run.text:
                    run.text = run.text.replace(clean_template, report_title)
                    return
            new_text = text.replace(clean_template, report_title)
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ''
            return
        break


# ============================================================
# Main Word generation function
# ============================================================

def _generate_word_from_template(template_file, title, results, template_structure=None, engine=None, selected_months=None, report_type='single'):
    """Generate Word report from template - placeholder fill, chart gen, table fill, summary gen"""
    from docx import Document

    doc = Document(template_file)

    # Step 1: Get summary data
    summary_data = _get_summary_data(engine, selected_months)
    month_str = selected_months[0] if selected_months else ""
    year = month_str[:4] if len(month_str) >= 6 else ""
    month_num = str(int(month_str[4:6])) if len(month_str) >= 6 else ""

    # Step 1.5: Update document title with actual report period
    if report_type == 'compare' and len(selected_months) >= 2:
        m1 = str(int(selected_months[0][4:6])) if len(selected_months[0]) >= 6 else "?"
        m2 = str(int(selected_months[1][4:6])) if len(selected_months[1]) >= 6 else "?"
        report_title = f"{m1}月 vs {m2}月 {title.replace('模板', '').strip()}"
    else:
        report_title = _build_report_title(title, year, month_num, selected_months)
    _update_document_title(doc, report_title, title)

    # Step 2: Replace placeholders (paragraphs + table cells)
    _fill_placeholders(doc, year, month_num, summary_data, engine=engine, selected_months=selected_months, report_type=report_type)

    # Step 3: Generate charts and replace description paragraphs
    _fill_section_charts(doc, engine, selected_months, report_type)

    # Step 4: Fill table data
    _fill_tables(doc, engine, selected_months, summary_data=summary_data, report_type=report_type)

    # Step 5: Generate summary & recommendations
    _fill_summary_section(doc, engine, selected_months, summary_data, report_type)

    # Save to memory
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ============================================================
# Route registration
# ============================================================

def register_template_export_routes(app, engine, protected=None):
    """Register template export routes"""
    protected = protected or _protected

    if engine:
        try:
            with engine.connect() as conn:
                conn.execute(sa_text(CREATE_TABLE_SQL))
                conn.commit()
            for sql in ALTER_TABLE_SQL:
                try:
                    with engine.connect() as conn:
                        conn.execute(sa_text(sql))
                        conn.commit()
                except Exception:
                    pass
        except Exception as e:
            print(f"report_templates table check failed: {e}")

    @app.route('/api/report-templates', methods=['GET'])
    @protected
    def list_report_templates():
        try:
            with engine.connect() as conn:
                result = conn.execute(sa_text("SELECT id, name, description, report_type, sections, template_file, created_at, updated_at FROM report_templates ORDER BY updated_at DESC"))
                templates = []
                for row in result.fetchall():
                    secs = json.loads(row[4]) if isinstance(row[4], str) else row[4]
                    t = {"id": row[0], "name": row[1], "description": row[2] or "", "report_type": row[3],
                         "sections": secs, "section_count": len(secs or []),
                         "template_file": row[5], "created_at": str(row[6]), "updated_at": str(row[7])}
                    templates.append(t)
                return jsonify({"templates": templates})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates', methods=['POST'])
    @protected
    def create_report_template():
        try:
            data = request.json
            name = data.get('name', '').strip()
            if not name:
                return jsonify({"error": "请输入模板名称"}), 400
            sections = data.get('sections', [])
            if not sections:
                return jsonify({"error": "请至少添加一个章节"}), 400

            with engine.connect() as conn:
                conn.execute(sa_text(
                    "INSERT INTO report_templates (name, description, report_type, sections, template_file, template_structure) "
                    "VALUES (:name, :desc, :rtype, :sections, :tfile, :tstruct)"
                ), {"name": name, "desc": data.get('description', ''), "rtype": data.get('report_type', 'single'),
                     "sections": json.dumps(sections, ensure_ascii=False),
                     "tfile": data.get('template_file'),
                     "tstruct": json.dumps(data.get('template_structure'), ensure_ascii=False) if data.get('template_structure') else None})
                conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates/<int:tid>', methods=['PUT'])
    @protected
    def update_report_template(tid):
        try:
            data = request.json
            with engine.connect() as conn:
                conn.execute(sa_text(
                    "UPDATE report_templates SET name=:name, description=:desc, report_type=:rtype, sections=:sections, "
                    "template_file=:tfile, template_structure=:tstruct WHERE id=:id"
                ), {"name": data.get('name', ''), "desc": data.get('description', ''),
                     "rtype": data.get('report_type', 'single'),
                     "sections": json.dumps(data.get('sections', []), ensure_ascii=False),
                     "tfile": data.get('template_file'),
                     "tstruct": json.dumps(data.get('template_structure'), ensure_ascii=False) if data.get('template_structure') else None,
                     "id": tid})
                conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates/<int:tid>', methods=['DELETE'])
    @protected
    def delete_report_template(tid):
        try:
            with engine.connect() as conn:
                conn.execute(sa_text("DELETE FROM report_templates WHERE id=:id"), {"id": tid})
                conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates/upload', methods=['POST'])
    @protected
    def upload_report_template():
        try:
            if 'file' not in request.files:
                return jsonify({"error": "请上传文件"}), 400
            file = request.files['file']
            if not file.filename.endswith('.docx'):
                return jsonify({"error": "仅支持.docx格式"}), 400

            from template_parser import parse_docx_template
            upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'uploads', 'templates')
            os.makedirs(upload_dir, exist_ok=True)
            import uuid
            uid = uuid.uuid4().hex[:32]
            original_name = os.path.splitext(file.filename)[0]
            save_name = f"{uid}_{original_name}.docx"
            save_path = os.path.join(upload_dir, save_name)
            file.save(save_path)

            structure = parse_docx_template(save_path)

            return jsonify({
                "success": True,
                "file_path": save_path,
                "original_filename": file.filename,
                "structure": structure
            })
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates/<int:tid>/execute', methods=['POST'])
    @protected
    def execute_report_template(tid):
        try:
            data = request.json or {}
            selected_months = data.get('months', [])

            with engine.connect() as conn:
                result = conn.execute(sa_text(
                    "SELECT name, report_type, sections FROM report_templates WHERE id=:id"
                ), {"id": tid})
                row = result.fetchone()
                if not row:
                    return jsonify({"error": "模板不存在"}), 404

            template_name, report_type, sections_json = row[0], row[1], row[2]
            sections = json.loads(sections_json) if isinstance(sections_json, str) else sections_json

            results = []
            for sec in sections:
                sec_title = sec.get('title', '')
                query_desc = sec.get('query', '') or infer_query_from_title(sec_title)
                if not query_desc:
                    continue
                try:
                    from analysis_routes import _build_analysis_prompt, _parse_llm_json, _validate_sql, _generate_summary, _call_llm as analysis_call_llm
                    schema_info = {}
                    data_samples = {}
                    try:
                        from analysis_routes import _get_schema_info, _get_data_samples
                        schema_info = _get_schema_info(engine)
                        data_samples = _get_data_samples(engine)
                    except Exception:
                        pass
                    messages = _build_analysis_prompt(query_desc, schema_info, [], data_samples, selected_months)
                    llm_resp = analysis_call_llm(messages)
                    if llm_resp:
                        spec = _parse_llm_json(llm_resp)
                        sql = spec.get('sql', '')
                        if sql and _validate_sql(sql):
                            with engine.connect() as conn:
                                df = pd.read_sql(sa_text(sql), conn)
                            answer = _generate_summary(df, spec)
                            chart = None
                            if not df.empty and spec.get('chart_type'):
                                chart = {
                                    "title": spec.get('title', sec_title),
                                    "chart_type": spec.get('chart_type', 'bar'),
                                    "x_field": spec.get('x_field', df.columns[0] if len(df.columns) > 0 else ''),
                                    "y_field": spec.get('y_field', df.columns[1] if len(df.columns) > 1 else ''),
                                    "data": json.loads(df.to_json(orient='records', force_ascii=False))
                                }
                            table_data = json.loads(df.to_json(orient='records', force_ascii=False)) if not df.empty else []
                            results.append({"title": sec_title, "answer": answer, "chart": chart, "table_data": table_data})
                        else:
                            results.append({"title": sec_title, "answer": "无法生成有效查询"})
                    else:
                        results.append({"title": sec_title, "answer": "LLM调用失败"})
                except Exception as e:
                    results.append({"title": sec_title, "answer": f"分析失败: {str(e)}"})

            return jsonify({"success": True, "name": template_name, "report_type": report_type, "results": results})
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500

    @app.route('/api/report-templates/<int:tid>/export', methods=['GET'])
    @protected
    def export_template(tid):
        try:
            months_param = request.args.get('months', '')
            selected_months = [m.strip() for m in months_param.split(',') if m.strip()] if months_param else []

            with engine.connect() as conn:
                result = conn.execute(sa_text(
                    "SELECT name, report_type, sections, template_file, template_structure "
                    "FROM report_templates WHERE id=:id"
                ), {"id": tid})
                row = result.fetchone()
                if not row:
                    return jsonify({"error": "模板不存在"}), 404

            template_name = row[0]
            report_type = row[1] or 'single'
            # Smart detection: if template name suggests comparison, treat as compare
            if report_type != 'compare' and ('对比' in template_name or 'compare' in template_name.lower()):
                report_type = 'compare'
            template_file = row[3]

            # Validate month selection for comparison reports
            if report_type == 'compare' and len(selected_months) < 2:
                return jsonify({"error": "对比分析报告至少需要选择2个月的数据"}), 400

            if not template_file or not os.path.exists(template_file):
                return jsonify({"error": "模板文件不存在"}), 404

            # Build proper report title with month info
            if report_type == 'compare' and len(selected_months) >= 2:
                m1 = str(int(selected_months[0][4:6])) if len(selected_months[0]) >= 6 else "?"
                m2 = str(int(selected_months[1][4:6])) if len(selected_months[1]) >= 6 else "?"
                report_title = f"{m1}月 vs {m2}月 {template_name.replace('模板', '').strip()}"
            else:
                month_str = selected_months[0] if selected_months else ""
                year = month_str[:4] if len(month_str) >= 6 else ""
                month_num = str(int(month_str[4:6])) if len(month_str) >= 6 else ""
                report_title = _build_report_title(template_name, year, month_num, selected_months)

            docx_bytes = _generate_word_from_template(
                template_file, template_name, [], None, engine, selected_months, report_type
            )

            return send_file(
                io.BytesIO(docx_bytes),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'{report_title}.docx',
            )
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500
