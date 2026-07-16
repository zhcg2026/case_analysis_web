# -*- coding: utf-8 -*-
"""知识库路由模块 - RAG知识库管理"""
import datetime
from flask import request, jsonify
from helpers import protected

def register_knowledge_routes(app, Session, engine, 
                              get_collection_stats, list_documents, 
                              delete_document, insert_document, 
                              search_similar, ask_question, init_rag):
    """注册知识库相关路由"""
    
    @app.route('/api/knowledge/stats', methods=['GET'])
    @protected
    def knowledge_stats():
        """获取知识库统计信息"""
        try:
            stats = get_collection_stats()
            return jsonify(stats), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/documents', methods=['GET'])
    @protected
    def knowledge_list_documents():
        """列出知识库中的所有文档"""
        try:
            docs = list_documents()
            return jsonify({'documents': docs, 'total': len(docs)}), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/documents/<doc_id>', methods=['DELETE'])
    @protected
    def knowledge_delete_document(doc_id):
        """删除指定文档"""
        try:
            result = delete_document(doc_id)
            if result['success']:
                return jsonify(result), 200
            else:
                return jsonify(result), 400
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/documents/batch-delete', methods=['POST'])
    @protected
    def knowledge_batch_delete():
        """批量删除文档"""
        try:
            data = request.get_json()
            doc_ids = data.get('doc_ids', [])

            if not doc_ids:
                return jsonify({'error': '请提供要删除的文档ID列表'}), 400

            success_count = 0
            failed_count = 0
            results = []

            for doc_id in doc_ids:
                result = delete_document(doc_id)
                if result['success']:
                    success_count += 1
                    results.append({'doc_id': doc_id, 'success': True})
                else:
                    failed_count += 1
                    results.append({'doc_id': doc_id, 'success': False, 'error': result.get('message', '删除失败')})

            return jsonify({
                'success': True,
                'message': f'成功删除 {success_count} 个文档，失败 {failed_count} 个',
                'success_count': success_count,
                'failed_count': failed_count,
                'results': results
            }), 200

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/documents/delete-all', methods=['POST'])
    @protected
    def knowledge_delete_all():
        """删除所有文档"""
        try:
            from rag import delete_all_documents
            result = delete_all_documents()
            return jsonify(result), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/upload', methods=['POST'])
    @protected
    def knowledge_upload_document():
        """上传文档到知识库"""
        try:
            # 检查是否有文件
            if 'file' not in request.files:
                # 也支持直接上传文本内容
                data = request.get_json()
                if data and 'content' in data:
                    doc_id = data.get('doc_id', f'doc_{datetime.datetime.now().strftime("%Y%m%d%H%M%S")}')
                    source = data.get('source', '手动输入')
                    content = data['content']
                    metadata = data.get('metadata', {})

                    result = insert_document(doc_id, content, source, metadata)
                    if result['success']:
                        return jsonify(result), 200
                    else:
                        return jsonify(result), 400

                return jsonify({'error': '没有文件或内容'}), 400

            file = request.files['file']
            if file.filename == '':
                return jsonify({'error': '没有选择文件'}), 400

            # 生成文档ID
            timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
            truncated_name = file.filename[:44] if len(file.filename) > 44 else file.filename
            doc_id = f'doc_{timestamp}_{truncated_name}'
            source = file.filename
            metadata = {
                'filename': file.filename,
                'uploaded_by': request.username,
                'upload_time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }

            # 读取文件内容
            content = ''
            filename_lower = file.filename.lower()

            if filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
                content = file.read().decode('utf-8', errors='ignore')
            elif filename_lower.endswith('.docx'):
                from docx import Document
                doc = Document(file)
                for para in doc.paragraphs:
                    content += para.text + '\n'
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + '\t'
                        content += '\n'
            else:
                content = file.read().decode('utf-8', errors='ignore')

            if not content.strip():
                return jsonify({'error': '文件内容为空'}), 400

            # 插入到向量库
            result = insert_document(doc_id, content, source, metadata)

            if result['success']:
                return jsonify({
                    'success': True,
                    'doc_id': doc_id,
                    'chunks': result['chunks'],
                    'message': f'文档上传成功，已分割为 {result["chunks"]} 个片段'
                }), 200
            else:
                return jsonify(result), 400

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/search', methods=['POST'])
    @protected
    def knowledge_search():
        """知识搜索"""
        try:
            data = request.get_json()
            query = data.get('query', '')
            top_k = data.get('top_k', 5)

            if not query:
                return jsonify({'error': '请提供搜索关键词'}), 400

            results = search_similar(query, top_k)
            return jsonify({'results': results, 'total': len(results)}), 200

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/ask', methods=['POST'])
    @protected
    def knowledge_ask():
        """知识问答"""
        try:
            data = request.get_json()
            question = data.get('question', '')
            top_k = data.get('top_k', 5)

            if not question:
                return jsonify({'error': '请提供问题'}), 400

            result = ask_question(question, top_k)
            return jsonify(result), 200

        except Exception as e:
            return jsonify({'error': str(e)}), 500

    @app.route('/api/knowledge/init', methods=['POST'])
    @protected
    def knowledge_init():
        """初始化知识库"""
        try:
            result = init_rag()
            return jsonify({'success': True, 'message': '知识库初始化成功'}), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500
