# -*- coding: utf-8 -*-
"""工具函数模块 - 从app.py提取的通用工具函数"""
import os
import re
import hashlib
import time
import datetime
import tempfile
import zipfile
import requests
import numpy as np
import pandas as pd
from functools import wraps
from flask import request, jsonify
from docx import Document

# 登录失败限制配置
LOGIN_ATTEMPTS = {}
MAX_LOGIN_ATTEMPTS = 5
LOCKOUT_DURATION = 300

def hash_password(password):
    import bcrypt
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def verify_password(password, hashed):
    import bcrypt
    if hashed.startswith("$2b$"):
        return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))
    else:
        return hashlib.sha256(password.encode()).hexdigest() == hashed

def is_strong_password(password):
    if len(password) < 8:
        return False, "密码长度至少8位"
    has_letter = any(c.isalpha() for c in password)
    has_digit = any(c.isdigit() for c in password)
    if not (has_letter and has_digit):
        return False, "密码必须包含字母和数字"
    return True, None

def check_login_attempts(username):
    if username not in LOGIN_ATTEMPTS:
        return True, None
    attempt = LOGIN_ATTEMPTS[username]
    if attempt["count"] >= MAX_LOGIN_ATTEMPTS:
        if time.time() < attempt["lock_until"]:
            remaining = int(attempt["lock_until"] - time.time())
            return False, f"账户已锁定，请{remaining}秒后再试"
        else:
            LOGIN_ATTEMPTS[username] = {"count": 0, "lock_until": 0}
    return True, None

def record_failed_login(username):
    if username not in LOGIN_ATTEMPTS:
        LOGIN_ATTEMPTS[username] = {"count": 0, "lock_until": 0}
    LOGIN_ATTEMPTS[username]["count"] += 1
    if LOGIN_ATTEMPTS[username]["count"] >= MAX_LOGIN_ATTEMPTS:
        LOGIN_ATTEMPTS[username]["lock_until"] = time.time() + LOCKOUT_DURATION

def clear_login_attempts(username):
    if username in LOGIN_ATTEMPTS:
        del LOGIN_ATTEMPTS[username]

SECRET_KEY = os.getenv("JWT_SECRET_KEY", "your-secret-key-for-jwt-token")
TOKEN_EXPIRATION = int(os.getenv("TOKEN_EXPIRATION_SECONDS", str(24 * 60 * 60)))

def generate_token(user_id, username, role):
    import jwt
    payload = {
        "user_id": user_id,
        "username": username,
        "role": role,
        "exp": datetime.datetime.utcnow() + datetime.timedelta(seconds=TOKEN_EXPIRATION)
    }
    return jwt.encode(payload, SECRET_KEY, algorithm="HS256")

def verify_token(token):
    import jwt
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        return payload
    except:
        return None

def protected(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
        if token.startswith("Bearer "):
            token = token[7:]
        payload = verify_token(token)
        if not payload:
            return jsonify({"error": "Invalid or expired token"}), 401
        request.user_id = payload["user_id"]
        request.username = payload["username"]
        request.role = payload["role"]
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
        if token.startswith("Bearer "):
            token = token[7:]
        payload = verify_token(token)
        if not payload:
            return jsonify({"error": "Invalid or expired token"}), 401
        if payload["role"] != "admin":
            return jsonify({"error": "Admin permission required"}), 403
        request.user_id = payload["user_id"]
        request.username = payload["username"]
        request.role = payload["role"]
        return f(*args, **kwargs)
    return decorated

def get_json_payload():
    data = request.get_json(silent=True)
    return data if isinstance(data, dict) else {}

def get_case_or_404(session, case_id, Case):
    case = session.query(Case).filter_by(id=case_id).first()
    if not case:
        return None, (jsonify({"error": "案件不存在"}), 404)
    return case, None

def generate_slug(text):
    slug = text.lower()
    slug = re.sub(r"\s+", "-", slug)
    slug = re.sub(r"[^\u4e00-\u9fa5a-z0-9-]", "", slug)
    slug = re.sub(r"-+", "-", slug)
    slug = slug.strip("-")
    if not slug:
        slug = hashlib.md5(text.encode()).hexdigest()[:8]
    return slug

def convert_nan_to_null(obj):
    if isinstance(obj, dict):
        return {key: convert_nan_to_null(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_nan_to_null(item) for item in obj]
    elif isinstance(obj, float) and np.isnan(obj):
        return None
    else:
        return obj

API_KEY = "58a51ac5-3b75-4c5e-85ac-1fb4ef652bd0"
API_URL = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"
MODEL = "doubao-seed-1-8-251228"
BAILIAN_GENERAL_API_KEY = "sk-8f9b17ffd00148868cdadcac65220930"
BAILIAN_GENERAL_API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
BAILIAN_GENERAL_MODEL = "qwen-plus"
BAILIAN_CHENGGUANTONG_API_KEY = "sk-9ee20f6ad5dd459aa8952e5ae979bead"
BAILIAN_CHENGGUANTONG_API_URL = "https://dashscope.aliyuncs.com/api/v1/apps/b608e4ed05c44c19bf7e71679c859689/completion"
API_CONNECT_TIMEOUT = 10
API_READ_TIMEOUT = 300
API_MAX_RETRIES = 3
API_RETRY_DELAY = 5

def call_llm_api(api_url, api_key, model, messages, max_tokens=3000, temperature=0.3, provider_name="LLM"):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
        "Accept": "application/json",
        "Connection": "keep-alive"
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens
    }
    retry_delay = API_RETRY_DELAY
    for attempt in range(API_MAX_RETRIES):
        try:
            print(f"[{provider_name}] 尝试调用 API ({attempt + 1}/{API_MAX_RETRIES})...")
            response = requests.post(
                api_url,
                headers=headers,
                json=payload,
                timeout=(API_CONNECT_TIMEOUT, API_READ_TIMEOUT),
                proxies={"http": None, "https": None}
            )
            print(f"[{provider_name}] 响应状态码: {response.status_code}")
            if response.status_code != 200:
                error_text = response.text[:500] if response.text else "无响应内容"
                print(f"[{provider_name}] HTTP 错误: {response.status_code}, 响应: {error_text}")
                if 400 <= response.status_code < 500:
                    return False, f"API 请求错误 ({response.status_code}): {error_text}"
                response.raise_for_status()
            result = response.json()
            if "choices" not in result or len(result["choices"]) == 0:
                print(f"[{provider_name}] 响应结构异常: {result}")
                return False, "API 响应格式异常: 缺少 choices 字段"
            if "message" not in result["choices"][0] or "content" not in result["choices"][0]["message"]:
                print(f"[{provider_name}] 响应结构异常")
                return False, "API 响应格式异常: 缺少 message.content 字段"
            content = result["choices"][0]["message"]["content"]
            print(f"[{provider_name}] API 调用成功, 响应长度: {len(content)}")
            return True, content
        except requests.exceptions.Timeout as e:
            print(f"[{provider_name}] 请求超时: {e}")
            if attempt < API_MAX_RETRIES - 1:
                print(f"[{provider_name}] {retry_delay}秒后重试...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                return False, f"API 调用超时，已重试 {API_MAX_RETRIES} 次"
        except requests.exceptions.ConnectionError as e:
            print(f"[{provider_name}] 连接错误: {e}")
            if attempt < API_MAX_RETRIES - 1:
                print(f"[{provider_name}] {retry_delay}秒后重试...")
                time.sleep(retry_delay)
                retry_delay *= 2
            else:
                return False, f"网络连接失败: {str(e)}"
        except requests.exceptions.RequestException as e:
            print(f"[{provider_name}] 请求异常: {e}")
            return False, f"API 请求失败: {str(e)}"
        except Exception as e:
            print(f"[{provider_name}] 未知异常: {e}")
            import traceback
            traceback.print_exc()
            return False, f"API 调用异常: {str(e)}"
    return False, "API 调用失败: 超过最大重试次数"

def clean_problem_description(text):
    if pd.isna(text) or text.strip() == "":
        return text
    text_str = str(text)
    text_str = re.sub(r"1[3-9]\d{9}", "", text_str)
    text_str = re.sub(r"0\d{2,3}-?\d{7,8}", "", text_str)
    text_str = re.sub(r"[\u4e00-\u9fa5][A-Za-z]·?[A-Za-z0-9]{4,6}", "", text_str)
    text_str = re.sub(r"[0-9A-Za-z]{4,}", "", text_str)
    text_str = re.sub(r"原转办编号：\d+", "", text_str)
    text_str = re.sub(r"[\u4e00-\u9fa5]{1,2}[先生|女士|小姐|同志]", "", text_str)
    text_str = re.sub(r"[0-9]+[单元|号楼|楼|室|房|号]", "", text_str)
    text_str = re.sub(r"\s+", " ", text_str).strip()
    return text_str

def desensitize_name(name):
    if pd.isna(name) or str(name).strip() == "":
        return name
    name_str = str(name).strip()
    if len(name_str) <= 1:
        return name_str
    return name_str[0] + "*" * (len(name_str) - 1)

def desensitize_phone(phone):
    if pd.isna(phone) or str(phone).strip() == "":
        return phone
    phone_str = str(phone).strip()
    if len(phone_str) <= 3:
        return phone_str
    return phone_str[:3] + "*" * 8

def desensitize_landline(landline):
    if pd.isna(landline) or str(landline).strip() == "":
        return landline
    landline_str = str(landline).strip()
    if len(landline_str) <= 4:
        return landline_str
    return landline_str[:-4] + "*" * 4

def desensitize_address(address):
    if pd.isna(address) or str(address).strip() == "":
        return address
    address_str = str(address).strip()
    parts = address_str.split(" ")
    if len(parts) <= 1:
        parts = re.split(r"[,，]", address_str)
    if len(parts) >= 3:
        return " ".join(parts[:3]) + " ****"
    elif len(parts) >= 2:
        return " ".join(parts[:2]) + " ****"
    else:
        return address_str[:4] + " ****"

def clean_and_desensitize_data(df, fields_config):
    result_df = df.copy()
    for field, field_types in fields_config.items():
        if field not in result_df.columns:
            continue
        if not isinstance(field_types, list):
            field_types = [field_types]
        for field_type in field_types:
            if field_type == "problem_description":
                result_df[field] = result_df[field].apply(clean_problem_description)
            elif field_type == "name":
                result_df[field] = result_df[field].apply(desensitize_name)
            elif field_type == "phone":
                result_df[field] = result_df[field].apply(desensitize_phone)
            elif field_type == "landline":
                result_df[field] = result_df[field].apply(desensitize_landline)
            elif field_type == "address":
                result_df[field] = result_df[field].apply(desensitize_address)
    return result_df

def extract_location_from_text(text):
    if pd.isna(text) or text.strip() == "":
        return "未提取到地址"
    parts = re.split(r"，|,|。|；|：", str(text).strip())
    stop_words = [
        "绿地内", "人行道", "非机动车道", "主干道", "垃圾", "经营", "乱放",
        "晾晒", "粪便", "摊点", "尘土", "满溢", "不洁", "摆乱放", "果皮箱外",
        "成袋垃圾", "动物粪便", "流动", "道路尘土", "外观不洁", "把式车辆",
        "机动车道","店外经营", "路面", "底盖", "小广告", "广告", "乱晾",
        "乱晒", "外墙", "线体", "车轮", "违规", "主次干道", "配电箱", "干枝", "户外"
    ]
    location_parts = []
    for part in parts:
        if any(word in part for word in stop_words):
            break
        if part.strip():
            location_parts.append(part.strip())
    if not location_parts:
        location_parts = [p.strip() for p in parts[:2] if p.strip()]
    return "，".join(location_parts) if location_parts else "未提取到地址"


def read_file_content(file):
    """读取文件内容，支持docx和xlsx文件"""
    import pandas as pd
    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension == '.docx':
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
            file.save(temp.name)
            temp_path = temp.name

        try:
            def extract_headers_footers(doc):
                texts = []
                try:
                    for section in doc.sections:
                        header = section.header
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text:
                                texts.append(f"页眉: {text}")
                        footer = section.footer
                        for para in footer.paragraphs:
                            text = para.text.strip()
                            if text:
                                texts.append(f"页脚: {text}")
                except Exception as e:
                    print(f"Error extracting headers/footers: {str(e)}")
                return texts

            doc = Document(temp_path)
            full_text = []

            header_footer_texts = extract_headers_footers(doc)
            if header_footer_texts:
                full_text.extend(header_footer_texts)

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_text.append(cell_text)
                    row_content = '\t'.join(row_text)
                    if row_content.strip():
                        full_text.append(row_content)

            try:
                with zipfile.ZipFile(temp_path, 'r') as zf:
                    if 'word/document.xml' in zf.namelist():
                        with zf.open('word/document.xml') as f:
                            xml_content = f.read().decode('utf-8')
                            text_content = re.sub('<[^<]+?>', '', xml_content)
                            text_content = text_content.strip()
                            if text_content and not full_text:
                                full_text.append(text_content)
            except Exception:
                pass

            content = '\n'.join(full_text)
            print(f"DOCX file processed: {len(content)} characters extracted")
            return content
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    elif file_extension == '.xlsx':
        df = pd.read_excel(file)
        content = []
        for index, row in df.iterrows():
            row_content = []
            for col in df.columns:
                if pd.notna(row[col]):
                    row_content.append(f"{col}: {row[col]}")
            if row_content:
                content.append(' | '.join(row_content))
        return '\n'.join(content)
    else:
        raise ValueError('Unsupported file type')
