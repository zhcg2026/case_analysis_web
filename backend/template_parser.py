# -*- coding: utf-8 -*-
"""Word 模板解析器 - 提取模板结构、图片位置等信息"""
import os
import re
import json
import shutil
from docx import Document
from docx.oxml.ns import qn


def _infer_chart_type(text):
    """根据文本描述推断图表类型"""
    t = text.lower()
    # 优先级：横向柱状图 > 饼状图/饼图 > 折线图 > 柱状图
    if "横向柱状图" in text or "水平柱状图" in text:
        return "horizontal_bar"
    if "饼状图" in text or "饼图" in text:
        return "pie"
    if "折线图" in text:
        return "line"
    if "柱状图" in text or "条形图" in text or "直方图" in text:
        return "bar"
    # 兜底：根据语义关键词再猜一次
    if "占比" in text or "比例" in text or "百分比" in text:
        return "pie"
    if "趋势" in text or "走势" in text or "变化" in text:
        return "line"
    return "bar"


def _extract_chart_name(text):
    """从图表描述文本中提取图表名称
    例如：'每日上报趋势（每日案件数量折线图）' -> '每日上报趋势'
          '24小时分布（分时段案件数量柱状图）' -> '24小时分布'
    """
    text = text.strip()
    # 去掉尾部排版说明，如"三图并排"、"两图并排"、"单图一列"
    text = re.sub(r"[（(]?[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行).*?[）)]?", "", text).strip()
    text = re.sub(r"[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行)", "", text).strip()
    # 匹配形如 '名称（...图）' 或 '名称(...图)'
    m = re.match(r"(.+?)[（(].*?[图)].*", text)
    if m:
        return m.group(1).strip()
    # 若含'图'字，取图字前面作为名称
    if "图" in text:
        idx = text.index("图")
        # 往前找一个合适的分隔位置（顿号、逗号、空格）
        prefix = text[:idx]
        for sep in ["、", "，", ",", " "]:
            if sep in prefix:
                return prefix.rsplit(sep, 1)[-1].strip()
        return prefix.strip()
    return text


def _split_chart_descriptions(text):
    """把一个段落里并列的多个图表描述拆成独立项
    例如：'每日上报趋势（折线图）、24小时分布（柱状图）、问题类型占比（饼状图）三图并排'
    -> ['每日上报趋势（折线图）', '24小时分布（柱状图）', '问题类型占比（饼状图）']
    """
    # 先去掉尾部"三图并排""两图并列"等说明
    text = re.sub(r"[（(]?[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行).*?[）)]?", "", text).strip()
    text = re.sub(r"[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行)", "", text).strip()
    # 按顿号/逗号拆分，保留括号内的内容不拆
    parts = []
    buf = ""
    depth = 0
    for ch in text:
        if ch in "（(":
            depth += 1
            buf += ch
        elif ch in "）)":
            depth -= 1
            buf += ch
        elif ch in "、，," and depth == 0:
            if buf.strip():
                parts.append(buf.strip())
            buf = ""
        else:
            buf += ch
    if buf.strip():
        parts.append(buf.strip())
    return parts


# 常见图表类型词，用于识别"纯类型词"（不是图表名称）
_TYPE_WORDS = ["横向柱状图", "水平柱状图", "柱状图", "条形图", "直方图",
               "饼状图", "饼图", "折线图", "对比图", "趋势图", "双轴图", "散点图"]


def _is_pure_type_word(s):
    """判断一段文字是否只是图表类型描述（如'横向柱状图'），而非图表名称"""
    t = (s or "").strip()
    if t in _TYPE_WORDS:
        return True
    cleaned = t
    for w in _TYPE_WORDS:
        cleaned = cleaned.replace(w, "")
    cleaned = re.sub(r"[图图表]+", "", cleaned)
    cleaned = cleaned.strip(" 、，,。；;（）()")
    return cleaned == ""


def _split_chart_specs(text):
    """从一段图表描述文本中解析出一组 (名称, 描述)。

    兼容三种写法：
      A) 带图号: '图1 问题大类数量对比：竖轴大类名称，横轴案件数...横向柱状图；图2 大类变化率：...；两图并列'
      B) 冒号无图号: '处置部门平均处置时长对比：竖轴处置部门，横轴平均时长...横向柱状图'
      C) 老并列写法（回退）: '每日上报趋势（折线图）、24小时分布（柱状图）三图并排'
                            或 '案件总量对比，5月案件数，6月案件数，柱状图'
    """
    text = (text or "").strip()
    # 去掉尾部"N图并排/并列/一列/一行"等排版说明（前面可能有空格、括号等）
    text = re.sub(r"[）)\s]*[（(]?[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行).*?[）)]?$", "", text).strip()
    text = re.sub(r"\s*[一二两三四五六七八九十\d]+图.*?(并排|并列|一列|一行)$", "", text).strip()
    if not text:
        return []

    # —— 格式A：存在"图N"标记（图+数字/中文数字，且前面是句首/分号/空格，排除"单图一列"等排版词） ——
    fig_pattern = re.compile(r"(?:^|[；;，,\s])图\s*([0-9一二三四五六七八九十]+)")
    fig_matches = list(fig_pattern.finditer(text))
    if fig_matches:
        specs = []
        for i, m in enumerate(fig_matches):
            start = m.end()
            end = fig_matches[i + 1].start() if i + 1 < len(fig_matches) else len(text)
            seg = text[start:end].strip().strip("；;。、")
            if not seg:
                continue
            ci = seg.find("：") if "：" in seg else seg.find(":")
            if ci >= 0:
                name = seg[:ci].strip()
                desc = seg[ci + 1:].strip()
            else:
                name = seg
                desc = ""
            if name and not _is_pure_type_word(name):
                specs.append((name, desc))
        if specs:
            return specs

    # —— 格式B：存在冒号且无图号 ——
    ci = text.find("：") if "：" in text else text.find(":")
    if ci > 0:
        name = text[:ci].strip()
        desc = text[ci + 1:].strip()
        if name and not _is_pure_type_word(name):
            return [(name, desc)]

    # —— 格式C：老并列写法，回退到原有逻辑 ——
    parts = _split_chart_descriptions(text)
    specs = []
    for p in parts:
        name = _extract_chart_name(p)
        # 用原始片段判断是否为纯类型词，避免旧逻辑把"柱状图"误提为"柱状"
        if name and not _is_pure_type_word(p):
            specs.append((name, p))
    return specs


def parse_docx_template(file_path):
    """
    解析 Word 模板，返回结构化信息
    
    Returns:
        dict: {
            "sections": [...],        # 章节列表
            "images": [...],          # 图片信息
            "tables": [...],          # 表格信息
            "summary_section": {...}  # 总结章节信息
        }
    """
    doc = Document(file_path)
    
    result = {
        "sections": [],
        "images": [],
        "tables": [],
        "summary_section": None
    }
    
    # 收集所有段落信息
    paragraphs_info = []
    for i, p in enumerate(doc.paragraphs):
        para_info = {
            "index": i,
            "style": p.style.name,
            "text": p.text,
            "has_image": _has_image(p),
            "image_info": _get_image_info(p) if _has_image(p) else None,
            "image_count": _count_images(p),
        }
        paragraphs_info.append(para_info)
    
    # 识别章节（从 Heading 1 样式）
    sections = []
    for i, para in enumerate(paragraphs_info):
        if para["style"] == "Heading 1":
            section = {
                "index": len(sections),
                "title": para["text"],
                "heading_paragraph_index": para["index"],
                "image_paragraph_index": None,
                "caption_paragraph_index": None,
                "table_index": None,
                "query": "",  # 兼容旧字段
                "chart_type": "bar",  # 兼容旧字段
                "charts": [],  # 新增：章节下的图表列表
            }
            
            # 收集本节范围内的图片段落，用于按顺序匹配图表
            section_image_paras = []
            
            # 向后查找图片、图注和图表描述
            next_heading_idx = len(paragraphs_info)
            for k in range(i + 1, len(paragraphs_info)):
                if paragraphs_info[k]["style"] == "Heading 1":
                    next_heading_idx = k
                    break
            
            for j in range(i + 1, next_heading_idx):
                next_para = paragraphs_info[j]
                
                # 跳过子标题（Heading 2 等）—— 遇到子标题也停止收集图表描述
                if next_para["style"] in ["Heading 2", "Heading 3"]:
                    continue
                
                # 找到图片，记录到本节图片列表
                if next_para["has_image"]:
                    section_image_paras.append(next_para)
                    result["images"].append({
                        "paragraph_index": next_para["index"],
                        "section_index": len(sections),
                        "image_info": next_para["image_info"]
                    })
                    if section["image_paragraph_index"] is None:
                        section["image_paragraph_index"] = next_para["index"]
                
                # 找到图注
                if next_para["style"] == "Caption" and section["caption_paragraph_index"] is None:
                    section["caption_paragraph_index"] = next_para["index"]
                
                # 识别图表描述：段落里出现"...图"
                if "图" in next_para["text"] or "chart" in next_para["text"].lower():
                    # 从描述段落中解析出每个图表的 (名称, 描述)
                    chart_specs = _split_chart_specs(next_para["text"])
                    for (chart_name, chart_desc) in chart_specs:
                        if not chart_name:
                            continue
                        # 从描述或名称中推断图表类型
                        inferred = _infer_chart_type(chart_desc or chart_name)
                        if not inferred:
                            continue
                        # 按顺序分配最近的图片段落
                        assigned_image_para = None
                        image_para_index_in_section = len(section["charts"])
                        if image_para_index_in_section < len(section_image_paras):
                            assigned_image_para = section_image_paras[image_para_index_in_section]["index"]
                        elif section_image_paras:
                            assigned_image_para = section_image_paras[-1]["index"]

                        section["charts"].append({
                            "name": chart_name,
                            "chart_type": inferred,
                            "query": "",
                            "description": chart_desc,
                            "image_paragraph_index": assigned_image_para,
                        })
            
            # 兼容旧结构：如果识别到 charts，用第一个图表填充旧字段
            if section["charts"]:
                section["chart_type"] = section["charts"][0]["chart_type"]
                section["query"] = section["charts"][0].get("query", "")
            
            sections.append(section)
    
    result["sections"] = sections
    
    # 识别表格位置
    for i, table in enumerate(doc.tables):
        table_pos = _get_table_paragraph_index(doc, i)
        # 找到表格前面最近的标题
        table_section_index = None
        for j, section in enumerate(sections):
            if section["heading_paragraph_index"] < len(doc.paragraphs):
                # 检查表格是否在这个章节范围内
                if j < len(sections) - 1:
                    next_section = sections[j + 1]
                    # 表格在当前章节标题之后、下一章节标题之前
                    if section["heading_paragraph_index"] <= table_pos < next_section["heading_paragraph_index"]:
                        table_section_index = j
                        break
                else:
                    # 最后一个章节：表格在其标题之后
                    if table_pos >= section["heading_paragraph_index"]:
                        table_section_index = j
                        break
        
        # 如果表格在第一个章节标题之前，关联到第一个章节
        if table_section_index is None and sections and table_pos < sections[0]["heading_paragraph_index"]:
            table_section_index = 0
        
        result["tables"].append({
            "table_index": i,
            "rows": len(table.rows),
            "cols": len(table.columns) if table.rows else 0,
            "section_index": table_section_index,
            "header": [cell.text[:30] for cell in table.rows[0].cells] if table.rows else []
        })
        
        # 关联表格到章节
        if table_section_index is not None and sections[table_section_index]["table_index"] is None:
            sections[table_section_index]["table_index"] = i
    
    # 识别总结章节
    for i, para in enumerate(paragraphs_info):
        if para["style"] == "Heading 1" and "总结" in para["text"]:
            summary = {
                "heading_paragraph_index": para["index"],
                "title": para["text"],
                "subsections": []
            }
            
            # 查找子章节
            for j in range(i + 1, len(paragraphs_info)):
                next_para = paragraphs_info[j]
                if next_para["style"] == "Heading 2":
                    summary["subsections"].append({
                        "title": next_para["text"],
                        "paragraph_index": next_para["index"]
                    })
                elif next_para["style"] == "Heading 1":
                    break
            
            result["summary_section"] = summary
            break
    
    return result


def extract_template_images(file_path, output_dir):
    """
    提取模板中的所有图片
    
    Args:
        file_path: Word 文件路径
        output_dir: 图片输出目录
    
    Returns:
        list: 图片信息列表
    """
    doc = Document(file_path)
    images = []
    
    os.makedirs(output_dir, exist_ok=True)
    
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            blob = rel.target_part.blob
            ext = rel.target_ref.split(".")[-1]
            filename = f"image_{i}.{ext}"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(blob)
            
            images.append({
                "rId": rel.rId,
                "filename": filename,
                "filepath": filepath,
                "size": len(blob),
                "extension": ext
            })
    
    return images


def copy_template_file(source_path, dest_dir):
    """
    复制模板文件到目标目录
    
    Args:
        source_path: 源文件路径
        dest_dir: 目标目录
    
    Returns:
        str: 复制后的文件路径
    """
    os.makedirs(dest_dir, exist_ok=True)
    dest_path = os.path.join(dest_dir, "template.docx")
    shutil.copy2(source_path, dest_path)
    return dest_path


def replace_image_in_doc(doc, paragraph_index, new_image_bytes, image_ext="png"):
    """
    替换文档中指定段落的图片
    
    Args:
        doc: python-docx Document 对象
        paragraph_index: 段落索引
        new_image_bytes: 新图片的字节数据
        image_ext: 图片扩展名
    
    Returns:
        bool: 是否成功替换
    """
    if paragraph_index >= len(doc.paragraphs):
        return False
    
    paragraph = doc.paragraphs[paragraph_index]
    
    # 查找段落中的图片元素
    drawings = paragraph._element.findall(f".//{qn('w:drawing')}")
    if not drawings:
        return False
    
    for drawing in drawings:
        # 查找图片引用
        blips = drawing.findall(f".//{qn('a:blip')}")
        for blip in blips:
            r_embed = blip.get(qn("r:embed"))
            if r_embed and r_embed in doc.part.rels:
                # 替换图片内容
                rel = doc.part.rels[r_embed]
                rel.target_part._blob = new_image_bytes
                return True
    
    return False


def replace_table_in_doc(doc, table_index, new_data, max_rows=20):
    """
    替换文档中指定索引的表格数据
    
    Args:
        doc: python-docx Document 对象
        table_index: 表格索引
        new_data: 新数据（列表字典）
        max_rows: 最大行数
    
    Returns:
        bool: 是否成功替换
    """
    if table_index >= len(doc.tables):
        return False
    
    table = doc.tables[table_index]
    if not new_data:
        return False
    
    # 获取列名
    cols = list(new_data[0].keys())
    
    # 清空现有表格（保留表头行）
    while len(table.rows) > 1:
        row = table.rows[-1]
        row._element.getparent().remove(row._element)
    
    # 更新表头
    header_row = table.rows[0]
    for j, col in enumerate(cols):
        if j < len(header_row.cells):
            cell = header_row.cells[j]
            cell.text = str(col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # 添加数据行
    for row_data in new_data[:max_rows]:
        row = table.add_row()
        for j, col in enumerate(cols):
            if j < len(row.cells):
                val = row_data.get(col, "")
                if isinstance(val, float):
                    row.cells[j].text = f"{val:,.2f}"
                else:
                    row.cells[j].text = str(val) if val else ""
    
    return True


def _count_images(paragraph):
    """统计段落中包含的图片数量"""
    drawings = paragraph._element.findall(f".//{qn('w:drawing')}")
    return len(drawings)


def _has_image(paragraph):
    """检查段落是否包含图片"""
    return _count_images(paragraph) > 0


def _get_image_info(paragraph):
    """获取段落中图片的信息"""
    drawings = paragraph._element.findall(f".//{qn('w:drawing')}")
    images = []
    
    for drawing in drawings:
        blips = drawing.findall(f".//{qn('a:blip')}")
        for blip in blips:
            r_embed = blip.get(qn("r:embed"))
            if r_embed:
                images.append({"rId": r_embed})
    
    return images[0] if images else None


def _get_table_paragraph_index(doc, table_index):
    """获取表格在文档中的精确段落位置（通过分析XML结构）"""
    from docx.oxml.ns import qn
    
    if table_index >= len(doc.tables):
        return len(doc.paragraphs)
    
    # 获取文档body的所有子元素
    body = doc.element.body
    
    # 计算每个段落的全局索引
    para_index = 0
    table_count = 0
    
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        if tag == 'p':  # 段落
            para_index += 1
        elif tag == 'tbl':  # 表格
            if table_count == table_index:
                # 表格在其前面的段落之后
                return para_index
            table_count += 1
    
    return para_index
